from __future__ import annotations

import unittest
from io import BytesIO

from docx import Document

from word_constructor.app import (
    _ai_correct_slot_values,
    _document_plain_text,
    _extract_header_footer_placeholder_occurrences,
    _extract_placeholder_occurrences,
    _openai_placeholder_payload,
    _raw_placeholder_matches_from_doc,
    _should_preserve_ai_corrected_value,
    fill_docx,
)


class AiContextExtractionTests(unittest.TestCase):
    def _fixture_doc(self) -> Document:
        doc = Document()
        doc.add_paragraph("Body intro [BodyName] before table.")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Cell owns [CellValue] only"
        table.cell(0, 1).text = "Adjacent cell must not leak"
        doc.add_paragraph("Body after table must not leak into cell context.")
        footer = doc.sections[0].footer
        footer.paragraphs[0].text = "Footer boilerplate [FooterValue] Страница 1 из 2"

        raw = BytesIO()
        doc.save(raw)
        raw.seek(0)
        return Document(raw)

    def test_occurrence_contexts_are_structurally_isolated(self) -> None:
        doc = self._fixture_doc()
        values = {
            "BodyName": "Иванов Иван Иванович",
            "CellValue": "табличное значение",
            "FooterValue": "footer value",
        }

        occurrences = _extract_placeholder_occurrences(doc, values)
        by_key = {item["placeholder"]: item for item in occurrences}

        self.assertEqual(by_key["BodyName"]["source_type"], "body_paragraph")
        self.assertEqual(by_key["BodyName"]["source_path"], "paragraph[0]")
        self.assertIn("Body intro [BodyName] before table.", by_key["BodyName"]["context_text"])

        cell = by_key["CellValue"]
        self.assertEqual(cell["source_type"], "table_cell")
        self.assertEqual(cell["source_path"], "table[0].row[0].cell[0]")
        self.assertIn("Cell owns [CellValue] only", cell["context_text"])
        self.assertNotIn("Adjacent cell", cell["context_text"])
        self.assertNotIn("Body after table", cell["context_text"])
        self.assertNotIn("Footer boilerplate", cell["context_text"])

        self.assertNotIn("FooterValue", by_key)
        footer_occurrences = _extract_header_footer_placeholder_occurrences(doc, values)
        self.assertEqual(len(footer_occurrences), 1)
        footer = footer_occurrences[0]
        self.assertEqual(footer["source_type"], "footer")
        self.assertEqual(footer["source_path"], "footer[section=0]")
        self.assertIn("Footer boilerplate [FooterValue]", footer["context_text"])
        self.assertNotIn("Body after table", footer["context_text"])


    def test_signature_table_name_is_excluded_from_ai_and_preserved(self) -> None:
        doc = Document()
        doc.add_paragraph("Просим рассмотреть заявление Джумабаеву Розу Багиткалиевну.")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "[SignerTitle]"
        table.cell(0, 1).text = "[SignerName]"

        values = {
            "SignerTitle": "члена Правления - заместителя председателя Правления",
            "SignerName": "Н. Джамышев",
        }
        occurrences = _extract_placeholder_occurrences(doc, values)
        self.assertEqual(len(occurrences), 2)
        by_key = {item["placeholder"]: item for item in occurrences}
        title_occurrence = by_key["SignerTitle"]
        name_occurrence = by_key["SignerName"]

        self.assertEqual(title_occurrence["source_type"], "table_cell")
        self.assertEqual(title_occurrence["source_path"], "table[0].row[0].cell[0]")
        self.assertTrue(title_occurrence["ai_excluded"])
        self.assertTrue(title_occurrence["signature_title_normalize"])

        self.assertEqual(name_occurrence["source_type"], "table_cell")
        self.assertEqual(name_occurrence["source_path"], "table[0].row[0].cell[1]")
        self.assertTrue(name_occurrence["ai_excluded"])
        self.assertEqual(name_occurrence["ai_exclusion_reason"], "signature_or_approval_table")
        self.assertIn("[SignerName]", name_occurrence["context_text"])
        self.assertNotIn("заявление", name_occurrence["context_text"])

        corrected_slots, occurrence_values = _ai_correct_slot_values(doc, values, "")
        self.assertEqual(corrected_slots, values)
        self.assertEqual(
            occurrence_values[("SignerTitle", 1)],
            "Члена правления - заместителя председателя правления",
        )
        self.assertEqual(occurrence_values[("SignerName", 1)], "Н. Джамышев")
        self.assertNotIn("\n", occurrence_values[("SignerName", 1)])
        self.assertEqual(" ".join(occurrence_values[("SignerName", 1)].split()), "Н. Джамышев")

        raw = BytesIO()
        doc.save(raw)
        output = Document(BytesIO(fill_docx(raw.getvalue(), corrected_slots, {}, slot_occurrence_values=occurrence_values)))
        self.assertEqual(output.tables[0].cell(0, 0).text, "Члена правления - заместителя председателя правления")
        self.assertEqual(output.tables[0].cell(0, 1).text, "Н. Джамышев")


    def test_short_from_table_cell_occurrence_is_sent_to_ai(self) -> None:
        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        table.cell(0, 0).text = ""
        table.cell(0, 1).text = "Директору департамента"
        table.cell(1, 0).text = ""
        table.cell(1, 1).text = "по управлению персоналом"
        table.cell(2, 0).text = ""
        table.cell(2, 1).text = "Сидоровой Е.Е."
        table.cell(3, 0).text = ""
        table.cell(3, 1).text = ""
        addressee_para = table.cell(3, 1).paragraphs[0]
        addressee_para.add_run("От ")
        addressee_para.add_run("[ФИО")
        addressee_para.add_run("Сотрудника]")
        doc.add_paragraph("Прошу предоставить [ФИОСотрудника] отпуск.")

        values = {"ФИОСотрудника": "Иванов Иван Иванович"}
        raw_matches = _raw_placeholder_matches_from_doc(doc, values)
        occurrences = _extract_placeholder_occurrences(doc, values)

        self.assertEqual(len(raw_matches), 2)
        self.assertEqual(len(occurrences), 2)
        self.assertEqual([item["occurrence_index"] for item in occurrences], [0, 1])
        self.assertEqual(occurrences[0]["source_type"], "table_cell")
        self.assertEqual(occurrences[0]["source_path"], "table[0].row[3].cell[1]")
        self.assertIn("От [ФИОСотрудника]", occurrences[0]["context_text"])
        self.assertFalse(occurrences[0]["ai_excluded"])
        self.assertEqual(occurrences[1]["source_type"], "body_paragraph")
        self.assertIn("Прошу предоставить [ФИОСотрудника]", occurrences[1]["context_text"])

        payload = _openai_placeholder_payload(values, {}, "", occurrences, _document_plain_text(doc))
        sent_occurrences = payload["messages"][1]["content"]
        self.assertIn('"occurrence_index": 0', sent_occurrences)
        self.assertIn('"occurrence_index": 1', sent_occurrences)
        self.assertIn('"source_path": "table[0].row[3].cell[1]"', sent_occurrences)

        corrected_slots, occurrence_values = _ai_correct_slot_values(doc, values, "")
        self.assertEqual(corrected_slots, values)
        self.assertEqual(occurrence_values[("ФИОСотрудника", 1)], "Иванова Ивана Ивановича")
        self.assertEqual(occurrence_values[("ФИОСотрудника", 2)], "Иванову Ивану Ивановичу")

        raw = BytesIO()
        doc.save(raw)
        output = Document(BytesIO(fill_docx(raw.getvalue(), corrected_slots, {}, slot_occurrence_values=occurrence_values)))
        self.assertEqual(output.tables[0].cell(3, 1).text, "От Иванова Ивана Ивановича")
        self.assertIn("Иванову Ивану Ивановичу", output.paragraphs[-1].text)

    def test_halyk_jumabayeva_occurrences_are_both_sent_to_ai(self) -> None:
        doc = Document()
        doc.add_paragraph("Принять [ФИОСотрудника] на должность главного менеджера управления учета брокерской деятельности АО Halyk Finance.")
        doc.add_paragraph("Основание: трудовой договор № [НомерДоговора] от [ДатаНачалаДоговора] года, заявление [ФИОСотрудника].")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "[ДолжностьСогласующего]"
        table.cell(0, 1).text = "[ФИОСогласующего]"

        values = {
            "ФИОСотрудника": "Джумабаева Роза Багиткалиевна",
            "НомерДоговора": "45/2026",
            "ДатаНачалаДоговора": "01.07.2026",
            "ДолжностьСогласующего": "члена Правления - заместителя председателя Правления",
            "ФИОСогласующего": "Ф.ОХОНОВ",
        }
        raw_matches = _raw_placeholder_matches_from_doc(doc, values)
        occurrences = _extract_placeholder_occurrences(doc, values)

        self.assertEqual(len(raw_matches), 6)
        self.assertEqual(len(occurrences), 6)
        employee_occurrences = [item for item in occurrences if item["placeholder"] == "ФИОСотрудника"]
        self.assertEqual(len(employee_occurrences), 2)
        self.assertEqual(employee_occurrences[0]["source_type"], "body_paragraph")
        self.assertEqual(employee_occurrences[0]["occurrence_index"], 0)
        self.assertIn("Принять [ФИОСотрудника]", employee_occurrences[0]["context_text"])
        self.assertEqual(employee_occurrences[1]["source_type"], "body_paragraph")
        self.assertEqual(employee_occurrences[1]["occurrence_index"], 1)
        self.assertIn("заявление [ФИОСотрудника]", employee_occurrences[1]["context_text"])

        payload = _openai_placeholder_payload(values, {}, "", occurrences, _document_plain_text(doc))
        sent_occurrences = payload["messages"][1]["content"]
        self.assertIn('"placeholder": "ФИОСотрудника", "occurrence_index": 0', sent_occurrences)
        self.assertIn('"placeholder": "ФИОСотрудника", "occurrence_index": 1', sent_occurrences)
        self.assertNotIn('"placeholder": "ДолжностьСогласующего"', sent_occurrences)
        self.assertNotIn('"placeholder": "ФИОСогласующего"', sent_occurrences)

        by_key = {item["placeholder"]: item for item in occurrences}
        self.assertTrue(by_key["ДолжностьСогласующего"]["ai_excluded"])
        self.assertTrue(by_key["ДолжностьСогласующего"]["signature_title_normalize"])
        self.assertTrue(by_key["ФИОСогласующего"]["ai_excluded"])

        corrected_slots, occurrence_values = _ai_correct_slot_values(doc, values, "")
        self.assertEqual(corrected_slots, values)
        self.assertEqual(
            occurrence_values[("ФИОСотрудника", 1)],
            "Джумабаеву Розу Багиткалиевну",
        )
        self.assertEqual(
            occurrence_values[("ФИОСотрудника", 2)],
            "Джумабаевой Розы Багиткалиевны",
        )
        self.assertEqual(occurrence_values[("НомерДоговора", 1)], "45/2026")
        self.assertEqual(occurrence_values[("ДатаНачалаДоговора", 1)], "01 июля 2026")
        self.assertEqual(occurrence_values[("ДолжностьСогласующего", 1)], "Члена правления - заместителя председателя правления")
        self.assertEqual(occurrence_values[("ФИОСогласующего", 1)], "Ф.Охонов")

        raw = BytesIO()
        doc.save(raw)
        output = Document(BytesIO(fill_docx(raw.getvalue(), corrected_slots, {}, slot_occurrence_values=occurrence_values)))
        self.assertIn("Принять Джумабаеву Розу Багиткалиевну", output.paragraphs[0].text)
        self.assertIn("№ 45/2026 от 01 июля 2026 года", output.paragraphs[1].text)
        self.assertIn("заявление Джумабаевой Розы Багиткалиевны", output.paragraphs[1].text)


    def test_kazakh_name_and_department_code_safeguards(self) -> None:
        doc = Document()
        doc.add_paragraph(
            "Принять [СсылкаСотрудникФизическоеЛицоФИО] на должность "
            "[РеквизитыСотрудникДолжностьНаименование] "
            "[РеквизитыСотрудникПодразделениеНаименование] АО «Halyk Finance»."
        )
        doc.add_paragraph(
            "Основание: трудовой договор № [НомерДоговора] от [ДатаНачалаДоговора] года, "
            "заявление [СсылкаСотрудникФизическоеЛицоФИО]."
        )
        values = {
            "СсылкаСотрудникФизическоеЛицоФИО": "Садык Ермек Жәнібекұлы",
            "РеквизитыСотрудникДолжностьНаименование": "hr бизнес-партнер",
            "РеквизитыСотрудникПодразделениеНаименование": "КОМПАС ВОАД",
            "НомерДоговора": "77/2026",
            "ДатаНачалаДоговора": "15.08.2026",
        }

        corrected_slots, occurrence_values = _ai_correct_slot_values(doc, values, "")
        self.assertEqual(
            occurrence_values[("СсылкаСотрудникФизическоеЛицоФИО", 1)],
            "Садыка Ермека Жәнібекұлы",
        )
        self.assertEqual(
            occurrence_values[("СсылкаСотрудникФизическоеЛицоФИО", 2)],
            "Садыка Ермека Жәнібекұлы",
        )
        self.assertEqual(occurrence_values[("НомерДоговора", 1)], "77/2026")
        self.assertEqual(occurrence_values[("ДатаНачалаДоговора", 1)], "15 августа 2026")
        self.assertEqual(occurrence_values[("РеквизитыСотрудникДолжностьНаименование", 1)], "HR бизнес-партнер")
        self.assertTrue(
            _should_preserve_ai_corrected_value(
                "РеквизитыСотрудникПодразделениеНаименование",
                "КОМПАС ВОАД",
                "компаса воад",
            )
        )

        raw = BytesIO()
        doc.save(raw)
        output = Document(BytesIO(fill_docx(raw.getvalue(), corrected_slots, {}, slot_occurrence_values=occurrence_values)))
        self.assertIn("Принять Садыка Ермека Жәнібекұлы", output.paragraphs[0].text)
        self.assertIn("HR бизнес-партнер КОМПАС ВОАД", output.paragraphs[0].text)
        self.assertIn("№ 77/2026 от 15 августа 2026 года", output.paragraphs[1].text)
        self.assertIn("заявление Садыка Ермека Жәнібекұлы", output.paragraphs[1].text)


    def test_consolidated_useai_regression_fixture(self) -> None:
        doc = Document()
        addressee = doc.add_table(rows=4, cols=2)
        addressee.cell(0, 1).text = "Директору департамента"
        addressee.cell(1, 1).text = "по управлению персоналом"
        addressee.cell(2, 1).text = "Сидоровой Е.Е."
        para = addressee.cell(3, 1).paragraphs[0]
        para.add_run("От ")
        para.add_run("[ФИО")
        para.add_run("Сотрудника]")
        doc.add_paragraph(
            "Принять [СсылкаСотрудникФизическоеЛицоФИО] на должность "
            "[РеквизитыСотрудникДолжностьНаименование] "
            "[РеквизитыСотрудникПодразделениеНаименование] АО «Halyk Finance»."
        )
        doc.add_paragraph("Прошу предоставить [ФИОСотрудника] отпуск.")
        doc.add_paragraph(
            "Основание: трудовой договор № [НомерДоговора] от [ДатаНачалаДоговора] года, "
            "заявление [СсылкаСотрудникФизическоеЛицоФИО]."
        )
        signature = doc.add_table(rows=1, cols=2)
        signature.cell(0, 0).text = "[ДолжностьСогласующего]"
        signature.cell(0, 1).text = "[ФИОСогласующего]"

        values = {
            "ФИОСотрудника": "Иванов Иван Иванович",
            "СсылкаСотрудникФизическоеЛицоФИО": "Садык Ермек Жәнібекұлы",
            "РеквизитыСотрудникДолжностьНаименование": "hr бизнес-партнер",
            "РеквизитыСотрудникПодразделениеНаименование": "Департамент ВОАД",
            "НомерДоговора": "77/2026",
            "ДатаНачалаДоговора": "15.08.2026",
            "ДолжностьСогласующего": "члена Правления - заместителя председателя Правления",
            "ФИОСогласующего": "Есжанова Зарина Серикалиевна",
        }

        raw_matches = _raw_placeholder_matches_from_doc(doc, values)
        occurrences = _extract_placeholder_occurrences(doc, values)
        self.assertEqual(len(raw_matches), len(occurrences))
        self.assertEqual(len(raw_matches), 10)

        dept = [item for item in occurrences if item["placeholder"] == "РеквизитыСотрудникПодразделениеНаименование"][0]
        self.assertTrue(dept["fixed_form"])
        self.assertTrue(dept["never_merge_with_adjacent_occurrence"])

        corrected_slots, occurrence_values = _ai_correct_slot_values(
            doc,
            values,
            "Номер договора нужно написать прописью",  # standing deterministic rules must still win
        )
        self.assertEqual(occurrence_values[("ФИОСотрудника", 1)], "Иванова Ивана Ивановича")
        self.assertEqual(occurrence_values[("ФИОСотрудника", 2)], "Иванову Ивану Ивановичу")
        self.assertEqual(occurrence_values[("СсылкаСотрудникФизическоеЛицоФИО", 1)], "Садыка Ермека Жәнібекұлы")
        self.assertEqual(occurrence_values[("СсылкаСотрудникФизическоеЛицоФИО", 2)], "Садыка Ермека Жәнібекұлы")
        self.assertEqual(occurrence_values[("НомерДоговора", 1)], "77/2026")
        self.assertEqual(occurrence_values[("ДатаНачалаДоговора", 1)], "15 августа 2026")
        self.assertEqual(occurrence_values[("РеквизитыСотрудникДолжностьНаименование", 1)], "HR бизнес-партнер")
        self.assertEqual(occurrence_values[("РеквизитыСотрудникПодразделениеНаименование", 1)], "Департамент ВОАД")
        self.assertNotIn("бизнес", occurrence_values[("РеквизитыСотрудникПодразделениеНаименование", 1)].lower())
        self.assertNotIn("ВОАД", occurrence_values[("РеквизитыСотрудникДолжностьНаименование", 1)])
        self.assertEqual(occurrence_values[("ДолжностьСогласующего", 1)], "Члена правления - заместителя председателя правления")
        self.assertEqual(occurrence_values[("ФИОСогласующего", 1)], "Есжанова З.С.")

        raw = BytesIO()
        doc.save(raw)
        output = Document(BytesIO(fill_docx(raw.getvalue(), corrected_slots, {}, slot_occurrence_values=occurrence_values)))
        self.assertEqual(output.tables[0].cell(3, 1).text, "От Иванова Ивана Ивановича")
        self.assertIn("Принять Садыка Ермека Жәнібекұлы", output.paragraphs[0].text)
        self.assertIn("HR бизнес-партнер Департамент ВОАД", output.paragraphs[0].text)
        self.assertIn("Иванову Ивану Ивановичу", output.paragraphs[1].text)
        self.assertIn("№ 77/2026 от 15 августа 2026 года", output.paragraphs[2].text)
        self.assertIn("заявление Садыка Ермека Жәнібекұлы", output.paragraphs[2].text)
        self.assertEqual(output.tables[1].cell(1 if len(output.tables) > 1 else 0, 1).text if False else output.tables[1].cell(0, 1).text, "Есжанова З.С.")

    def test_full_document_text_excludes_header_footer_and_labels_units(self) -> None:
        doc = self._fixture_doc()
        text = _document_plain_text(doc)

        self.assertIn("[body_paragraph paragraph[0]]", text)
        self.assertIn("[table_cell table[0].row[0].cell[0]]", text)
        self.assertIn("[table_cell table[0].row[0].cell[1]]", text)
        self.assertNotIn("Footer boilerplate", text)


if __name__ == "__main__":
    unittest.main()
