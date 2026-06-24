from __future__ import annotations

from io import BytesIO

from docx import Document

from word_constructor.app import _replace_values_from_json_payload, fill_docx


def _template_bytes() -> bytes:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Начисления"
    table.cell(0, 1).text = "Размер"
    table.cell(1, 0).text = "[Начисления.НоваяСтрока.Начисления_Начисление]"
    table.cell(1, 1).text = "[Начисления.НоваяСтрока.Начисления_Размер]"
    raw = BytesIO()
    doc.save(raw)
    return raw.getvalue()


def test_object_array_table_expands_nachisleniya_rows() -> None:
    rows = [
        {
            "Начисления_Начисление": "Оплата по окладу (по часам)",
            "Начисления_Размер": "742\u00a0000",
        },
        {"Начисления_Начисление": "Надбавка", "Начисления_Размер": "50000"},
    ]

    output = Document(
        BytesIO(fill_docx(_template_bytes(), {}, {}, table_object_params={"Начисления": rows}))
    )

    assert [[cell.text for cell in row.cells] for row in output.tables[0].rows] == [
        ["Начисления", "Размер"],
        ["Оплата по окладу (по часам)", "742\u00a0000"],
        ["Надбавка", "50000"],
    ]


def test_json_replace_payload_accepts_top_level_table_values() -> None:
    payload = {
        "filename": "template.docx",
        "content_base64": "ignored-in-this-unit-test",
        "UseAI": True,
        "PromtAI": "ignored control field",
        "ФИО": "Иванов И.И.",
        "Начисления": [
            {
                "Начисления_Начисление": "Оплата по окладу (по часам)",
                "Начисления_Размер": "742\u00a0000",
            },
        ],
    }

    values = _replace_values_from_json_payload(payload)

    assert values == {
        "ФИО": "Иванов И.И.",
        "Начисления": [
            {
                "Начисления_Начисление": "Оплата по окладу (по часам)",
                "Начисления_Размер": "742\u00a0000",
            },
        ],
    }
