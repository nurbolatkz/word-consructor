from __future__ import annotations

import base64
import json
import os
import tempfile
import time
import unittest
from io import BytesIO
from unittest.mock import patch

from docx import Document

from word_constructor.admin_storage import load_review_items
from word_constructor.ai_correction.claude_checker_and_summarizer import claude_correct_and_review
from app import create_app


class _FakeResponse:
    status = 200

    def __init__(self, body: dict):
        self.body = json.dumps(body, ensure_ascii=False).encode("utf-8")

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False

    def read(self) -> bytes:
        return self.body


def _docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class Stage2ReviewPipelineTests(unittest.TestCase):
    def test_claude_correct_and_review_reports_changes_and_note(self) -> None:
        gpt_response = {"РеквизитыРуководительФИО": "Есжанова З.С."}
        fake_body = {
            "corrected_values": {"РеквизитыРуководительФИО": "Есжанова З.С."},
            "review_summary": {
                "had_issues": False,
                "changes_from_gpt": [],
                "note": "No corrections needed — GPT's output matched.",
            },
        }

        with patch("word_constructor.ai_correction.claude_checker_and_summarizer._client") as client_factory:
            client = client_factory.return_value
            client.messages.create.return_value = type(
                "Resp",
                (),
                {"content": [type("Block", (), {"type": "text", "text": json.dumps(fake_body, ensure_ascii=False)})()]},
            )()
            result = claude_correct_and_review(
                {"template": "[РеквизитыРуководительФИО]", "placeholders": {"РеквизитыРуководительФИО": "Есжанова Зарина Серикалиевна"}},
                gpt_response,
            )

        self.assertEqual(result["corrected_values"]["РеквизитыРуководительФИО"], "Есжанова З.С.")
        self.assertFalse(result["review_summary"]["had_issues"])
        self.assertEqual(result["review_summary"]["changes_from_gpt"], [])
        self.assertTrue(result["review_summary"]["note"])


    def test_background_review_enqueue_is_non_blocking(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            env = {
                "ADMIN_REVIEW_DB_PATH": f"{tmp}/admin.sqlite3",
                "AI_CORRECTION_LOG_DIR": f"{tmp}/logs",
                "OPENAI_API_KEY": "test-key",
                "ANTHROPIC_API_KEY": "test-key",
            }
            payload = {
                "filename": "signature.docx",
                "content_base64": base64.b64encode(_docx_bytes("[РеквизитыРуководительФИО]")).decode("ascii"),
                "params": {"РеквизитыРуководительФИО": "Есжанова Зарина Серикалиевна"},
            }
            gpt_body = {
                "choices": [{"message": {"content": json.dumps({"РеквизитыРуководительФИО": "Есжанова З.С.", "_review_needed": False}, ensure_ascii=False)}}],
                "usage": {"total_tokens": 10},
            }
            claude_body = {
                "corrected_values": {"РеквизитыРуководительФИО": "Есжанова Зарина Серикалиевна"},
                "review_summary": {
                    "had_issues": True,
                    "changes_from_gpt": [
                        {
                            "placeholder": "РеквизитыРуководительФИО",
                            "gpt_value": "Есжанова З.С.",
                            "claude_value": "Есжанова Зарина Серикалиевна",
                            "reason": "GPT abbreviates FIO in a signature label.",
                        }
                    ],
                    "note": "Claude corrected the signature label.",
                },
            }

            def fake_anthropic_client():
                class Messages:
                    @staticmethod
                    def create(**kwargs):
                        return type(
                            "Resp",
                            (),
                            {"content": [type("Block", (), {"type": "text", "text": json.dumps(claude_body, ensure_ascii=False)})()]},
                        )()
                return type("Client", (), {"messages": Messages()})()

            app = create_app()
            client = app.test_client()
            with patch.dict(os.environ, env), patch("word_constructor.ai_correction.openai_client.urlopen", return_value=_FakeResponse(gpt_body)), patch("word_constructor.ai_correction.claude_checker_and_summarizer._client", side_effect=fake_anthropic_client), patch("word_constructor.admin_views._persist_background_review_log", side_effect=lambda payload: time.sleep(0.35)):
                t0 = time.perf_counter()
                resp = client.post("/services/word-constructor/api/replace", json=payload)
                elapsed = time.perf_counter() - t0

            self.assertLess(elapsed, 0.5)
            self.assertEqual(resp.status_code, 200)


    def test_sync_replace_endpoint_uses_claude_value_and_logs_background_review(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            env = {
                "ADMIN_REVIEW_DB_PATH": f"{tmp}/admin.sqlite3",
                "AI_CORRECTION_LOG_DIR": f"{tmp}/logs",
                "OPENAI_API_KEY": "test-key",
                "ANTHROPIC_API_KEY": "test-key",
            }
            payload = {
                "filename": "signature.docx",
                "content_base64": base64.b64encode(_docx_bytes("[РеквизитыРуководительФИО]")).decode("ascii"),
                "params": {"РеквизитыРуководительФИО": "Есжанова Зарина Серикалиевна"},
            }
            gpt_body = {
                "choices": [{"message": {"content": json.dumps({"РеквизитыРуководительФИО": "Есжанова З.С.", "_review_needed": False}, ensure_ascii=False)}}],
                "usage": {"total_tokens": 10},
            }
            claude_body = {
                "corrected_values": {"РеквизитыРуководительФИО": "Есжанова Зарина Серикалиевна"},
                "review_summary": {
                    "had_issues": True,
                    "changes_from_gpt": [
                        {
                            "placeholder": "РеквизитыРуководительФИО",
                            "gpt_value": "Есжанова З.С.",
                            "claude_value": "Есжанова Зарина Серикалиевна",
                            "reason": "GPT abbreviates FIO in a signature label.",
                        }
                    ],
                    "note": "Claude corrected the signature label.",
                },
            }

            def fake_anthropic_client():
                class Messages:
                    @staticmethod
                    def create(**kwargs):
                        return type(
                            "Resp",
                            (),
                            {"content": [type("Block", (), {"type": "text", "text": json.dumps(claude_body, ensure_ascii=False)})()]},
                        )()
                return type("Client", (), {"messages": Messages()})()

            app = create_app()
            client = app.test_client()
            t0 = time.perf_counter()
            with patch.dict(os.environ, env), patch("word_constructor.ai_correction.openai_client.urlopen", return_value=_FakeResponse(gpt_body)), patch("word_constructor.ai_correction.claude_checker_and_summarizer._client", side_effect=fake_anthropic_client):
                resp = client.post("/services/word-constructor/api/replace", json=payload)
                elapsed = time.perf_counter() - t0

                self.assertLess(elapsed, 1.0)
                self.assertEqual(resp.status_code, 200)
                result_doc = Document(BytesIO(resp.data))
                self.assertIn("Есжанова Зарина Серикалиевна", result_doc.paragraphs[0].text)
                items = []
                for _ in range(40):
                    items = load_review_items()
                    if items:
                        break
                    time.sleep(0.05)

                self.assertTrue(items)
                self.assertEqual(items[0]["document_name"], "signature.docx")
                self.assertTrue(items[0]["log_key"].startswith("replace-"))
                self.assertEqual(items[0]["corrections"][0]["placeholder"], "РеквизитыРуководительФИО")
                self.assertTrue(items[0]["checker_result"]["review_summary"]["had_issues"])


if __name__ == "__main__":
    unittest.main()
