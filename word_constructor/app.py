from __future__ import annotations

import base64
import json
import logging
import os


def _safe_b64decode(value: str) -> bytes:
    """
    Decode base64 tolerating 1C quirks:
      - MIME line-breaks (already stripped by caller, but double-safe)
      - URL-safe chars (- and _ instead of + and /)
      - Missing = padding
      - data-URI prefix
    """
    s = (value or "").strip()
    if ";base64," in s:
        s = s.split(";base64,", 1)[1]
    s = "".join(s.split())          # remove any remaining whitespace
    s = s.replace("-", "+").replace("_", "/")   # URL-safe → standard
    s += "=" * (-len(s) % 4)        # fix padding
    return base64.b64decode(s)
import copy
import hashlib
import hmac
import re
import secrets
import shutil
import threading
import time
import uuid
import zipfile
from datetime import datetime, timezone
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any, Optional
from urllib.parse import urlparse, urlunparse
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

import jwt
from docx import Document
from flask import Blueprint, abort, g, jsonify, redirect, render_template, request, send_file, session

from word_constructor.transforms import apply_transform, get_transforms

from word_constructor.ai_correction.deterministic import (
    format_ru_date_no_year_word as _format_ru_date_no_year_word_new,
    normalize_common_business_abbreviations as _normalize_common_business_abbreviations_new,
    normalize_signature_name as _normalize_signature_name_new,
    normalize_signature_title as _normalize_signature_title_new,
    should_preserve_ai_corrected_value as _should_preserve_ai_corrected_value_new,
)
from word_constructor.ai_correction.extraction import (
    PLACEHOLDER_RE as _AI_PLACEHOLDER_RE,
    cell_text as _ai_cell_text,
    context_snippet as _ai_context_snippet,
    document_full_text as _ai_document_full_text,
    document_plain_text as _ai_document_plain_text,
    document_placeholder_scan_text as _ai_document_placeholder_scan_text,
    extract_header_footer_placeholder_occurrences as _ai_extract_header_footer_placeholder_occurrences,
    extract_placeholder_contexts as _ai_extract_placeholder_contexts,
    extract_placeholder_occurrences as _ai_extract_placeholder_occurrences,
    iter_text_units as _ai_iter_text_units,
    match_key as _ai_match_key,
    raw_placeholder_matches_from_doc as _ai_raw_placeholder_matches_from_doc,
    sanity_check_occurrence_counts as _ai_sanity_check_occurrence_counts,
)
from word_constructor.ai_correction.claude_checker import claude_available as _ai_claude_available
from word_constructor.ai_correction.claude_checker_and_summarizer import claude_correct_occurrences as _ai_claude_correct_occurrences
from word_constructor.ai_correction.openai_client import parse_openai_chat_content as _ai_parse_openai_chat_content
from word_constructor.ai_correction.pipeline import correct_slot_values as _ai_pipeline_correct_slot_values
from word_constructor.ai_correction.pipeline import startup_health as ai_correction_startup_health
from word_constructor.admin_views import enqueue_background_review_log
from word_constructor.admin_storage import load_review_items as _load_review_items

logger = logging.getLogger(__name__)

word_constructor = Blueprint(
    "word_constructor",
    __name__,
    template_folder="templates",
)

STORAGE_DIR = Path("/tmp/kazuni_word_constructor")
STORAGE_DIR.mkdir(parents=True, exist_ok=True)
CLIENT_STORE_PATH = Path(os.environ.get("CLIENT_STORE_PATH", "/tmp/kazuni_word_constructor_clients.json"))
_CLIENT_STORE_LOCK = threading.Lock()
DEFAULT_SESSION_TTL_SECONDS = 35 * 60
SESSION_TTL_SECONDS = max(
    int(os.environ.get("SESSION_TTL_SECONDS", str(DEFAULT_SESSION_TTL_SECONDS))),
    60,
)
SESSION_TB_TTL_SECONDS = max(
    int(os.environ.get("TEMPLATE_BUILDER_SESSION_TTL_SECONDS", str(SESSION_TTL_SECONDS))),
    60,
)
FORCESAVE_WAIT_SECONDS = 8.0

# Matches both {{key}} and {{Table.row.N}} / {{Table.N.M}} (dot-notation cell refs)
# and [Key] (native 1C format).
# Group 1 = curly-brace key (may contain dots), Group 2 = square-bracket key.
_PLACEHOLDER_RE = re.compile(r"\{\{([^{}\n\r]{1,120})\}\}|\[([^\[\]\n\r]{1,120})\]")
AI_PLACEHOLDER_CONTEXT_CHARS = max(
    int(os.environ.get("AI_PLACEHOLDER_CONTEXT_CHARS", "240")),
    40,
)
AI_PLACEHOLDER_MAX_SNIPPETS = max(
    int(os.environ.get("AI_PLACEHOLDER_MAX_SNIPPETS", "5")),
    1,
)
OPENAI_PLACEHOLDER_TIMEOUT_SECONDS = max(
    float(os.environ.get("OPENAI_PLACEHOLDER_TIMEOUT_SECONDS", "8")),
    1.0,
)



def _match_key(m: re.Match) -> str:
    """Return the placeholder key regardless of which format matched."""
    return _ai_match_key(m)


def _resolve_table_cell(key: str, table_params: dict) -> str | None:
    """
    Resolve dot-notation table references:
      TableName.R.C   → cell at row R, col C (0-indexed; row 0 = header)
      TableName.row.R → comma-joined values of row R
      TableName.col.C → comma-joined values of column C
    Returns None if key is not a valid dot-notation table reference.
    """
    parts = key.split(".", 2)
    if len(parts) < 3:
        return None
    table_name, sub1, sub2 = parts
    if table_name not in table_params:
        return None
    rows = table_params[table_name]
    if not rows:
        return ""
    try:
        if sub1 == "row":
            r = int(sub2)
            return ", ".join(str(c) for c in rows[r]) if 0 <= r < len(rows) else ""
        elif sub1 == "col":
            c = int(sub2)
            return ", ".join(str(rows[r][c]) for r in range(len(rows)) if c < len(rows[r]))
        else:
            r, c = int(sub1), int(sub2)
            return str(rows[r][c]) if r < len(rows) and c < len(rows[r]) else ""
    except (ValueError, IndexError):
        return ""


# ---------------------------------------------------------------------------
# Admin clients and token auth
# ---------------------------------------------------------------------------

def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _hash_token(token: str) -> str:
    return hashlib.sha256(token.encode("utf-8")).hexdigest()


def _empty_client_store() -> dict[str, Any]:
    return {"clients": [], "admin": {}}


def _read_client_store() -> dict[str, Any]:
    if not CLIENT_STORE_PATH.exists():
        return _empty_client_store()
    try:
        raw = json.loads(CLIENT_STORE_PATH.read_text(encoding="utf-8"))
    except Exception:
        return _empty_client_store()
    if not isinstance(raw, dict) or not isinstance(raw.get("clients"), list):
        return _empty_client_store()
    if not isinstance(raw.get("admin"), dict):
        raw["admin"] = {}
    return raw


def _write_client_store(store: dict[str, Any]) -> None:
    CLIENT_STORE_PATH.parent.mkdir(parents=True, exist_ok=True)
    tmp = CLIENT_STORE_PATH.with_suffix(".tmp")
    tmp.write_text(json.dumps(store, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(CLIENT_STORE_PATH)


def _parse_admin_expires(value: str) -> str | None:
    raw = (value or "").strip()
    if not raw:
        return None
    dt = datetime.fromisoformat(raw)
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc).replace(microsecond=0).isoformat()


def _parse_iso(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        dt = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc)


def _client_is_expired(client: dict[str, Any], now: datetime | None = None) -> bool:
    expires_at = _parse_iso(client.get("expires_at"))
    if expires_at is None:
        return False
    return (now or datetime.now(timezone.utc)) > expires_at


def _client_public(client: dict[str, Any]) -> dict[str, Any]:
    stats = client.get("stats") if isinstance(client.get("stats"), dict) else {}
    return {
        "id": client.get("id", ""),
        "name": client.get("name", ""),
        "created_at": client.get("created_at", ""),
        "expires_at": client.get("expires_at"),
        "active": bool(client.get("active", True)),
        "expired": _client_is_expired(client),
        "stats": {
            "calls": int(stats.get("calls", 0) or 0),
            "input_bytes": int(stats.get("input_bytes", 0) or 0),
            "output_bytes": int(stats.get("output_bytes", 0) or 0),
            "last_call_at": stats.get("last_call_at"),
            "last_path": stats.get("last_path"),
        },
    }


def _extract_bearer_token() -> str:
    auth = request.headers.get("Authorization", "")
    if auth.lower().startswith("bearer "):
        return auth.split(None, 1)[1].strip()
    return (request.headers.get("X-Client-Token") or request.args.get("token") or "").strip()


def request_has_client_token() -> bool:
    return bool(_extract_bearer_token())


def client_api_index_response():
    client, error = _authenticate_api_client()
    if error is not None:
        return error
    g.api_client_id = client.get("id") if client else None
    g.api_client_name = client.get("name") if client else None
    return jsonify({
        "status": "ok",
        "service": "word-constructor",
        "client": {
            "id": client.get("id", "") if client else "",
            "name": client.get("name", "") if client else "",
        },
        "endpoints": {
            "replace": "/services/word-constructor/api/1c/replace",
            "replace_edit": "/services/word-constructor/api/1c/replace-edit",
            "template_builder_bridge": "/services/word-constructor/api/1c/template-builder/bridge",
            "word_base64_to_pdf": "/services/word-constructor/api/1c/converter/word-base64-to-pdf/",
            "sign_document": "/sign_document/api/1c/requests",
        },
    })


def _find_client_by_token(token: str) -> dict[str, Any] | None:
    if not token:
        return None
    token_hash = _hash_token(token)
    with _CLIENT_STORE_LOCK:
        store = _read_client_store()
        for client in store["clients"]:
            stored_hash = str(client.get("token_hash", ""))
            if stored_hash and hmac.compare_digest(stored_hash, token_hash):
                return client
    return None


def _authenticate_api_client() -> tuple[dict[str, Any] | None, tuple[Any, int] | None]:
    client = _find_client_by_token(_extract_bearer_token())
    if client is None:
        return None, (jsonify({"error": "Missing or invalid client token"}), 401)
    if not client.get("active", True):
        return None, (jsonify({"error": "Client token is disabled"}), 403)
    if _client_is_expired(client):
        return None, (jsonify({"error": "Client token expired"}), 403)
    return client, None


def _session_client_id_from_path(path: str) -> str | None:
    m = re.match(r"^/services/word-constructor/api/template-builder/([^/]+)/(?:status|download)$", path)
    if not m:
        return None
    meta = _read_meta(m.group(1))
    if not meta or meta.get("type") != "template_builder":
        return None
    return meta.get("client_id")


def _client_api_needs_token(path: str) -> bool:
    if path.startswith("/services/word-constructor/api/1c/"):
        return True
    return _session_client_id_from_path(path) is not None


def _record_client_usage(client_id: str, response) -> None:
    in_bytes = int(request.content_length or 0)
    out_bytes = response.calculate_content_length()
    out_bytes = int(out_bytes or 0)
    with _CLIENT_STORE_LOCK:
        store = _read_client_store()
        for client in store["clients"]:
            if client.get("id") != client_id:
                continue
            stats = client.setdefault("stats", {})
            stats["calls"] = int(stats.get("calls", 0) or 0) + 1
            stats["input_bytes"] = int(stats.get("input_bytes", 0) or 0) + in_bytes
            stats["output_bytes"] = int(stats.get("output_bytes", 0) or 0) + out_bytes
            stats["last_call_at"] = _utc_now_iso()
            stats["last_path"] = request.path
            _write_client_store(store)
            break


def _admin_logged_in() -> bool:
    return session.get("admin_logged_in") is True


def _admin_credentials_ok(username: str, password: str) -> bool:
    expected_user = os.environ.get("ADMIN_USERNAME", "admin")
    if not hmac.compare_digest(username, expected_user):
        return False

    with _CLIENT_STORE_LOCK:
        store = _read_client_store()
        stored_hash = str(store.get("admin", {}).get("password_hash", "") or "")

    if stored_hash:
        return hmac.compare_digest(stored_hash, _hash_token(password))

    expected_pass = os.environ.get("ADMIN_PASSWORD", "admin")
    return hmac.compare_digest(password, expected_pass)


def _set_admin_password(password: str) -> None:
    with _CLIENT_STORE_LOCK:
        store = _read_client_store()
        admin = store.setdefault("admin", {})
        admin["password_hash"] = _hash_token(password)
        admin["password_changed_at"] = _utc_now_iso()
        _write_client_store(store)


@word_constructor.before_request
def require_client_token():
    path = request.path
    if path.startswith("/services/word-constructor/admin"):
        return None
    if not _client_api_needs_token(path):
        return None

    client, error = _authenticate_api_client()
    if error is not None:
        return error
    required_client_id = _session_client_id_from_path(path)
    if required_client_id and client and client.get("id") != required_client_id:
        return jsonify({"error": "Token is not allowed for this session"}), 403
    g.api_client_id = client.get("id") if client else None
    g.api_client_name = client.get("name") if client else None
    return None


@word_constructor.after_request
def record_client_stats(response):
    client_id = getattr(g, "api_client_id", None)
    if client_id:
        _record_client_usage(client_id, response)
    return response

# ---------------------------------------------------------------------------
# Session helpers
# ---------------------------------------------------------------------------

def _session_dir(session_id: str) -> Path:
    return STORAGE_DIR / session_id


def _read_meta(session_id: str) -> dict | None:
    path = _session_dir(session_id) / "meta.json"
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def _google_client_id() -> str:
    secrets_path = Path("/opt/kazuni_doc_editor/google-api-secrets.json")
    if not secrets_path.exists():
        return ""
    try:
        raw = json.loads(secrets_path.read_text(encoding="utf-8"))
    except Exception:
        return ""
    web = raw.get("web") if isinstance(raw, dict) else {}
    if isinstance(web, dict):
        return str(web.get("client_id", "") or "")
    return ""


def _write_meta(session_id: str, meta: dict) -> None:
    path = _session_dir(session_id) / "meta.json"
    path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")


def _is_expired(meta: dict) -> bool:
    return time.time() > meta.get("expires_at", 0)


def _session_template_path(session_id: str) -> Path:
    return _session_dir(session_id) / "template.docx"


def _meta_string_params(meta: dict) -> dict[str, str]:
    raw = meta.get("params", {})
    if isinstance(raw, dict):
        return {str(k): str(v) for k, v in raw.items()}
    return {}


def public_base_url(req) -> str:
    proto = req.headers.get("X-Forwarded-Proto", req.scheme)
    host = req.headers.get("X-Forwarded-Host") or req.headers.get("Host") or req.host
    return f"{proto}://{host}".rstrip("/")


def public_ws_base_url(req) -> str:
    proto = "wss" if req.headers.get("X-Forwarded-Proto", req.scheme) == "https" else "ws"
    host = req.headers.get("X-Forwarded-Host") or req.headers.get("Host") or req.host
    return f"{proto}://{host}".rstrip("/")


def _send_file_compat(path_or_file, *, as_attachment=False, download_name=None, mimetype=None):
    kwargs = {
        "as_attachment": as_attachment,
        "mimetype": mimetype,
        "conditional": False,
        "etag": False,
    }
    if download_name is not None:
        kwargs["download_name"] = download_name
    return send_file(path_or_file, **kwargs)


def _onlyoffice_internal_base_url() -> str:
    return os.environ.get("ONLYOFFICE_INTERNAL_BASE_URL", "http://host.docker.internal:8016").rstrip("/")


def _onlyoffice_service_base_url() -> str:
    return os.environ.get("ONLYOFFICE_SERVICE_BASE_URL", "http://127.0.0.1:8020").rstrip("/")


def _onlyoffice_public_base_url() -> str:
    return os.environ.get("ONLYOFFICE_PUBLIC_BASE_URL", "/onlyoffice").rstrip("/")


def _onlyoffice_api_url() -> str:
    return f"{_onlyoffice_public_base_url()}/web-apps/apps/api/documents/api.js"


def _onlyoffice_jwt_secret() -> str:
    return os.environ.get("ONLYOFFICE_JWT_SECRET", "kazuni-onlyoffice-secret")


def _onlyoffice_document_type(extension: str) -> str:
    if extension in {"xls", "xlsx", "ods", "csv"}:
        return "cell"
    if extension in {"ppt", "pptx", "odp"}:
        return "slide"
    return "word"


def _normalize_callback_download_url(url: str) -> tuple[str, str | None]:
    parsed = urlparse(url)
    if parsed.hostname in {"127.0.0.1", "localhost"}:
        return url, None

    service_base = urlparse(_onlyoffice_service_base_url())
    if service_base.scheme and service_base.netloc:
        rewritten = parsed._replace(
            scheme=service_base.scheme,
            netloc=service_base.netloc,
        )
        return urlunparse(rewritten), None

    return url, None


def _download_remote_file(url: str, host_header: str | None = None) -> bytes:
    req = Request(url, headers={"User-Agent": "kazuni-word-constructor/1.0"})
    if host_header:
        req.add_header("Host", host_header)
    with urlopen(req, timeout=60) as response:
        return response.read()


def _builder_internal_url(session_id: str, suffix: str) -> str:
    return (
        f"{_onlyoffice_internal_base_url()}/services/word-constructor/"
        f"api/template-builder/{session_id}/{suffix.lstrip('/')}"
    )


def _onlyoffice_command_urls(key: str) -> list[str]:
    base = _onlyoffice_service_base_url()
    return [
        f"{base}/command?shardkey={key}",
        f"{base}/coauthoring/CommandService.ashx",
    ]


def _onlyoffice_converter_urls() -> list[str]:
    base = _onlyoffice_service_base_url()
    return [
        f"{base}/ConvertService.ashx",
    ]


def _post_json(url: str, payload: dict[str, Any], timeout: int = 20) -> dict[str, Any]:
    req = Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Accept": "application/json",
            "User-Agent": "kazuni-word-constructor/1.0",
        },
        method="POST",
    )
    with urlopen(req, timeout=timeout) as response:
        raw = response.read().decode("utf-8")
    return json.loads(raw or "{}")


def _converter_source_path(conversion_id: str) -> Path:
    return _session_dir(conversion_id) / "source.docx"


def _converter_internal_url(conversion_id: str) -> str:
    return (
        f"{_onlyoffice_internal_base_url()}/services/word-constructor/"
        f"api/converter/{conversion_id}/source"
    )


def _convert_docx_to_pdf_with_onlyoffice(filename: str, document_bytes: bytes) -> bytes:
    conversion_id = str(uuid.uuid4())
    sdir = _session_dir(conversion_id)
    sdir.mkdir(parents=True, exist_ok=True)
    expires_at = time.time() + 5 * 60
    meta = {
        "id": conversion_id,
        "type": "conversion",
        "filename": filename,
        "expires_at": expires_at,
        "expires_at_iso": datetime.fromtimestamp(expires_at, timezone.utc).isoformat(),
    }
    _write_meta(conversion_id, meta)
    _converter_source_path(conversion_id).write_bytes(document_bytes)

    key = hashlib.sha256(document_bytes + conversion_id.encode("utf-8")).hexdigest()
    payload = {
        "async": False,
        "filetype": "docx",
        "key": key,
        "outputtype": "pdf",
        "title": filename,
        "url": _converter_internal_url(conversion_id),
    }
    payload["token"] = jwt.encode(payload, _onlyoffice_jwt_secret(), algorithm="HS256")

    last_error: Exception | None = None
    try:
        for url in _onlyoffice_converter_urls():
            try:
                result = _post_json(url, payload, timeout=60)
            except (HTTPError, URLError, TimeoutError, ValueError) as exc:
                last_error = exc
                continue

            error_code = int(result.get("error", 0) or 0)
            if error_code != 0:
                raise RuntimeError(f"ONLYOFFICE conversion failed with error {error_code}: {result}")
            file_url = result.get("fileUrl") or result.get("fileurl")
            if not file_url:
                raise RuntimeError(f"ONLYOFFICE conversion response has no fileUrl: {result}")
            download_url, host_header = _normalize_callback_download_url(str(file_url))
            return _download_remote_file(download_url, host_header)
        raise RuntimeError(f"Cannot reach ONLYOFFICE conversion service: {last_error}")
    finally:
        shutil.rmtree(sdir, ignore_errors=True)


def _builder_editor_key(session_id: str, path: Path) -> str:
    return f"{session_id}-{int(path.stat().st_mtime)}-{path.stat().st_size}"


def _builder_forcesave(session_id: str, key: str) -> dict[str, Any]:
    payload = {"c": "forcesave", "key": key}
    payload["token"] = jwt.encode(payload, _onlyoffice_jwt_secret(), algorithm="HS256")
    last_error: Exception | None = None
    for url in _onlyoffice_command_urls(key):
        try:
            return _post_json(url, payload)
        except (HTTPError, URLError, TimeoutError, ValueError) as exc:
            last_error = exc
            continue
    raise RuntimeError(f"Cannot reach ONLYOFFICE command service: {last_error}")


def _wait_for_builder_save(session_id: str, previous_saved_at: float | int | None) -> bool:
    deadline = time.time() + FORCESAVE_WAIT_SECONDS
    baseline = float(previous_saved_at or 0)
    while time.time() < deadline:
        meta = _read_meta(session_id)
        if meta is None:
            return False
        current_saved_at = float(meta.get("last_saved_at") or 0)
        if current_saved_at > baseline:
            return True
        time.sleep(0.25)
    return False


# ---------------------------------------------------------------------------
# Background cleanup
# ---------------------------------------------------------------------------

def _cleanup_loop() -> None:
    while True:
        time.sleep(60)
        now = time.time()
        for d in STORAGE_DIR.iterdir():
            if not d.is_dir():
                continue
            meta_file = d / "meta.json"
            try:
                if meta_file.exists():
                    meta = json.loads(meta_file.read_text(encoding="utf-8"))
                    if now > meta.get("expires_at", 0):
                        shutil.rmtree(d, ignore_errors=True)
                elif (now - d.stat().st_mtime) > SESSION_TTL_SECONDS:
                    shutil.rmtree(d, ignore_errors=True)
            except Exception:
                pass


threading.Thread(target=_cleanup_loop, daemon=True).start()


# ---------------------------------------------------------------------------
# docx utilities
# ---------------------------------------------------------------------------

def _para_full_text(para) -> str:
    """Merge all runs so placeholders split across runs are found correctly."""
    return "".join(r.text for r in para.runs)


def _extract_placeholder_keys(doc: Document) -> list[str]:
    keys: set[str] = set()
    for para in doc.paragraphs:
        for m in _PLACEHOLDER_RE.finditer(_para_full_text(para)):
            keys.add(_match_key(m))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for m in _PLACEHOLDER_RE.finditer(_para_full_text(para)):
                        keys.add(_match_key(m))
    return sorted(keys)


def _truthy_request_value(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "y", "on"}
    return False


def _parse_ai_replace_options() -> tuple[bool, str]:
    if request.is_json:
        payload = request.get_json(silent=True) or {}
        use_ai = _truthy_request_value(payload.get("UseAI", payload.get("use_ai", False)))
        prompt = str(payload.get("PromtAI", payload.get("PromptAI", payload.get("prompt_ai", ""))) or "")
        logger.info(
            "replace-edit AI options parsed: is_json=%s use_ai=%s prompt_present=%s payload_keys=%s",
            request.is_json,
            use_ai,
            bool(prompt.strip()),
            sorted(str(key) for key in payload.keys()),
        )
        return use_ai, prompt
    use_ai = _truthy_request_value(request.form.get("UseAI") or request.form.get("use_ai"))
    prompt = request.form.get("PromtAI") or request.form.get("PromptAI") or request.form.get("prompt_ai") or ""
    logger.info(
        "replace-edit AI options parsed: is_json=%s use_ai=%s prompt_present=%s form_keys=%s",
        request.is_json,
        use_ai,
        bool(str(prompt).strip()),
        sorted(str(key) for key in request.form.keys()),
    )
    return use_ai, str(prompt)


def _placeholder_context_snippet(text: str, match: re.Match, window: int) -> str:
    return _ai_context_snippet(text, match, window)


def _cell_text(cell) -> str:
    return _ai_cell_text(cell)


def _iter_docx_structural_text_units(doc: Document, include_headers_footers: bool = False):
    for unit in _ai_iter_text_units(doc, include_headers_footers=include_headers_footers):
        yield {
            "source_type": unit.source_type,
            "source_path": unit.source_path,
            "text": unit.text,
            "table_index": unit.table_index,
            "row_index": unit.row_index,
            "cell_index": unit.cell_index,
            "row_cell_texts": list(unit.row_cell_texts),
        }

def _iter_docx_text_blocks(doc: Document):
    for unit in _iter_docx_structural_text_units(doc):
        yield unit["text"]


def _document_plain_text(doc: Document) -> str:
    return _ai_document_plain_text(doc)

def _document_placeholder_scan_text(doc: Document) -> str:
    return _ai_document_placeholder_scan_text(doc)

def _extract_placeholder_contexts(doc: Document, slot_values: dict[str, str]) -> dict[str, list[str]]:
    return _ai_extract_placeholder_contexts(doc, slot_values, AI_PLACEHOLDER_MAX_SNIPPETS)


_SIGNATURE_TITLE_RE = re.compile(
    r"\b(?:член[а-я]*|правлени[яею]|председател[яьюе]?|заместител[яьюе]?|директор[а-я]*|"
    r"руководител[яьюе]?|начальник[а-я]*|исполнительн[а-я]+|генеральн[а-я]+)\b",
    re.IGNORECASE,
)
_INITIAL_SURNAME_RE = re.compile(r"^[А-ЯЁA-Z]\.?\s*[А-ЯЁA-Z][А-ЯЁа-яёA-Za-z\-]+$")
_FULL_NAME_RE = re.compile(r"^[А-ЯЁ][А-ЯЁа-яё\-]+(?:\s+[А-ЯЁ][А-ЯЁа-яё\-]+){1,3}$")
_VERB_HINT_RE = re.compile(
    r"\b(?:прошу|предоставить|назначить|уволить|перевести|согласовать|утвердить|"
    r"является|составил|подписал|обязать|направить|принять)\b",
    re.IGNORECASE,
)
_SIGNATURE_KEY_RE = re.compile(
    r"(?:подпис|соглас|утверд|руковод|директор|председател|заместител|sign|signer)",
    re.IGNORECASE,
)
_SIGNATURE_TITLE_KEY_RE = re.compile(
    r"(?:должност|позици|руковод|директор|председател|заместител|title|position)",
    re.IGNORECASE,
)


_RU_MONTHS_GENT = [
    "", "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def _format_ru_date_no_year_word(value: str) -> str:
    match = re.fullmatch(r"\s*(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{2}|\d{4})\s*", value or "")
    if not match:
        return value
    day, month, year = match.groups()
    try:
        month_idx = int(month)
    except ValueError:
        return value
    if not 1 <= month_idx <= 12:
        return value
    if len(year) == 2:
        year = f"20{year}"
    return f"{int(day):02d} {_RU_MONTHS_GENT[month_idx]} {year}"


def _fix_common_feminine_surname_case(original: str, declined: str, case: str) -> str:
    original_words = (original or "").split()
    declined_words = (declined or "").split()
    if len(original_words) < 2 or len(original_words) != len(declined_words):
        return declined

    surname = original_words[0]
    lower = surname.lower()
    replacement = None
    if lower.endswith(("ова", "ева", "ина")):
        stem = surname[:-1]
        if case == "accs":
            replacement = stem + "у"
        elif case in {"gent", "datv", "loct", "ablt"}:
            replacement = stem + "ой"
    elif lower.endswith("ая"):
        stem = surname[:-2]
        if case == "accs":
            replacement = stem + "ую"
        elif case in {"gent", "datv", "loct", "ablt"}:
            replacement = stem + "ой"

    if not replacement:
        return declined
    declined_words[0] = replacement
    return " ".join(declined_words)


def _preserve_kazakh_patronymic_suffixes(original: str, declined: str) -> str:
    original_words = (original or "").split()
    declined_words = (declined or "").split()
    if len(original_words) != len(declined_words):
        return declined
    for idx, word in enumerate(original_words):
        if word.lower().endswith(("ұлы", "улы", "қызы", "кизы")):
            declined_words[idx] = word
    return " ".join(declined_words)


def _is_code_like_token(token: str) -> bool:
    cleaned = token.strip(".,;:()[]{}«»\"'")
    return bool(2 <= len(cleaned) <= 12 and re.search(r"[A-ZА-ЯЁҰҚІҒӘҺӨҮ]", cleaned) and cleaned.upper() == cleaned)


def _normalize_common_business_abbreviations(value: str) -> str:
    replacements = {"hr": "HR", "it": "IT", "pr": "PR", "ceo": "CEO", "cfo": "CFO", "cto": "CTO"}

    def repl(match: re.Match) -> str:
        return replacements.get(match.group(0).lower(), match.group(0))

    return re.sub(r"\b(?:hr|it|pr|ceo|cfo|cto)\b", repl, value or "", flags=re.IGNORECASE)


def _is_title_or_department_key(key: str) -> bool:
    lower_key = key.lower()
    return any(part in lower_key for part in ("должност", "позици", "подраздел", "департамент", "отдел", "управлен", "title", "position", "department", "division"))


def _should_preserve_ai_corrected_value(key: str, original: str, corrected: str) -> bool:
    lower_key = key.lower()
    if not corrected or corrected == original:
        return False
    if any(part in lower_key for part in ("подраздел", "департамент", "отдел", "управлен", "department", "division")):
        original_tokens = original.split()
        if any(_is_code_like_token(token) for token in original_tokens):
            corrected_tokens = corrected.split()
            for token in original_tokens:
                if _is_code_like_token(token) and token not in corrected_tokens:
                    return True
    return False


def _looks_like_signature_name_or_label(value: str) -> bool:
    cleaned = re.sub(r"\s+", " ", value or "").strip()
    if not cleaned:
        return False
    if _INITIAL_SURNAME_RE.match(cleaned):
        return True
    return bool(_FULL_NAME_RE.match(cleaned))


def _looks_like_signature_title(value: str) -> bool:
    cleaned = re.sub(r"\s+", " ", value or "").strip()
    return bool(cleaned and _SIGNATURE_TITLE_RE.search(cleaned) and not _looks_like_signature_name_or_label(cleaned))


def _normalize_signature_title(value: str) -> str:
    cleaned = re.sub(r"\s+", " ", value or "").strip()
    if not cleaned:
        return value
    lowered = cleaned.lower()
    return lowered[:1].upper() + lowered[1:]


def _normalize_signature_name(value: str) -> str:
    cleaned = re.sub(r"\s+", " ", value or "").strip()
    if not cleaned:
        return value

    def fix_surname(match: re.Match) -> str:
        prefix = match.group("prefix")
        surname = match.group("surname")
        return prefix + surname[:1].upper() + surname[1:].lower()

    normalized = re.sub(
        r"(?P<prefix>(?:[А-ЯЁA-Z]\.\s*){1,3})(?P<surname>[А-ЯЁA-Z]{2,}(?:-[А-ЯЁA-Z]{2,})*)\b",
        fix_surname,
        cleaned,
    )

    words = normalized.split()
    if 1 <= len(words) <= 4 and any(word.isupper() and len(word) > 1 for word in words):
        fixed_words = []
        for word in words:
            if re.fullmatch(r"(?:[А-ЯЁA-Z]\.){1,3}", word):
                fixed_words.append(word.upper())
            elif word.isupper() and len(word) > 1:
                fixed_words.append(word[:1].upper() + word[1:].lower())
            else:
                fixed_words.append(word)
        normalized = " ".join(fixed_words)
    return normalized


def _is_signature_or_approval_table_cell(unit: dict[str, Any], key: str, value: str) -> bool:
    if unit.get("source_type") != "table_cell":
        return False
    text = re.sub(r"\s+", " ", str(unit.get("text") or "")).strip()
    row_texts = [re.sub(r"\s+", " ", str(item or "")).strip() for item in unit.get("row_cell_texts") or []]
    row_joined = " | ".join(row_texts)
    if _VERB_HINT_RE.search(text):
        return False

    placeholder_only_or_short = bool(_is_sole_placeholder(text)) or len(text) <= 120
    row_has_signature_title = bool(
        _SIGNATURE_TITLE_RE.search(row_joined)
        or _SIGNATURE_TITLE_RE.search(value)
        or _SIGNATURE_TITLE_KEY_RE.search(row_joined)
    )
    key_or_value_is_signatory = bool(_SIGNATURE_KEY_RE.search(key)) or _looks_like_signature_name_or_label(value)
    return placeholder_only_or_short and row_has_signature_title and key_or_value_is_signatory


def _raw_placeholder_matches_from_doc(doc: Document, slot_values: dict[str, Any]) -> list[dict[str, Any]]:
    return _ai_raw_placeholder_matches_from_doc(doc, slot_values)


def _log_placeholder_occurrence_count_check(
    doc: Document,
    slot_values: dict[str, Any],
    occurrences: list[dict[str, Any]],
    log_key: str | None = None,
) -> None:
    raw_matches = _raw_placeholder_matches_from_doc(doc, slot_values)
    scan_text = _document_placeholder_scan_text(doc)
    wanted = set(slot_values)
    full_text_raw_count = sum(1 for match in _PLACEHOLDER_RE.finditer(scan_text) if _match_key(match) in wanted)
    raw_count = len(raw_matches)
    occurrence_count = len(occurrences)
    message_data = {
        "use_ai_log_key": log_key,
        "full_text_raw_match_count": full_text_raw_count,
        "raw_match_count": raw_count,
        "occurrence_count": occurrence_count,
        "raw_matches": raw_matches,
        "occurrences": [
            {
                "placeholder": item.get("placeholder", item.get("key")),
                "occurrence_index": item.get("occurrence_index"),
                "source_type": item.get("source_type"),
                "source_path": item.get("source_path"),
                "ai_excluded": item.get("ai_excluded"),
                "ai_exclusion_reason": item.get("ai_exclusion_reason"),
                "context_text": item.get("context_text"),
            }
            for item in occurrences
        ],
    }
    if full_text_raw_count != occurrence_count or raw_count != occurrence_count:
        logger.error(
            "UseAI placeholder occurrence mismatch before OpenAI: found %s raw [Placeholder] regex matches in full extracted document text, found %s source-aware raw matches, but added %s occurrences: %s",
            full_text_raw_count,
            raw_count,
            occurrence_count,
            message_data,
        )
    else:
        logger.debug(
            "UseAI placeholder occurrence count check passed: found %s raw [Placeholder] regex matches in full extracted document text and added %s occurrences: %s",
            full_text_raw_count,
            occurrence_count,
            message_data,
        )


def _extract_placeholder_occurrences(doc: Document, slot_values: dict[str, str]) -> list[dict[str, Any]]:
    return _ai_extract_placeholder_occurrences(doc, slot_values)


def _extract_header_footer_placeholder_occurrences(doc: Document, slot_values: dict[str, str]) -> list[dict[str, Any]]:
    return _ai_extract_header_footer_placeholder_occurrences(doc, slot_values)


AI_PLACEHOLDER_EXAMPLES = {
    "examples": [
        {
            "id": "signature_table_bug",
            "description": "Name in a signature/approval table wrongly declined to dative, inheriting case from an unrelated body paragraph. Also fixing capitalization of 'правления' (lowercase except first word of the phrase).",
            "template_placeholders": {
                "ДолжностьСогласующего": "члена Правления - заместителя председателя Правления",
                "ФИОСогласующего": "Н. Джамышев",
            },
            "context_snippet": "Table row: [ДолжностьСогласующего] | [ФИОСогласующего]",
            "current_wrong_output": "члена Правления - заместителя председателя Правления    Н. Джамышеву",
            "expected_output": "Члена правления - заместителя председателя правления    Н. Джамышев",
            "expected_case": "nominative",
            "capitalization_rule": "Capitalize only the first letter of the whole job-title phrase; 'правления' is a common noun here (not part of a proper name like 'Совет директоров'/'Правление' as an organization name on its own) and should be lowercase in subsequent occurrences within the same phrase.",
            "reason": "Signature block identifies who signed; it is a label, not a grammatical object of a verb in a sentence. Job-title phrases in such tables follow standard sentence-case capitalization: first word capitalized, the rest lowercase unless they are themselves proper nouns.",
        },
        {
            "id": "table_cell_correct_declension_contrast",
            "description": "Contrast case — a name inside a table cell that SHOULD be declined because the row's own text contains a grammatical role (genitive after 'заявления').",
            "template_placeholders": {
                "ФИОСотрудника": "Иванов Иван Иванович",
                "Должность": "ведущий специалист",
            },
            "context_snippet": "Table column 'Основание': на основании заявления [ФИОСотрудника]",
            "current_wrong_output": None,
            "expected_output": "на основании заявления Иванова Ивана Ивановича",
            "expected_case": "genitive",
            "reason": "The cell's own text contains a preposition+noun ('заявления') that grammatically governs the case of the name. This is not the signature-table case — correction should still apply here.",
        },
        {
            "id": "body_paragraph_case_bleed",
            "description": "A second occurrence of the same name in a different body paragraph incorrectly reuses the dative/accusative case from an earlier paragraph instead of being analyzed independently.",
            "template_placeholders": {
                "ФИОСотрудника": "Джумабаева Роза Багиткалиевна",
            },
            "context_snippet_occurrence_0": "Принять [ФИОСотрудника] на должность главного менеджера управления учета брокерской деятельности АО «Halyk Finance»",
            "context_snippet_occurrence_1": "Основание: трудовой договор № [НомерДоговора] от [ДатаНачалаДоговора] года, заявление [ФИОСотрудника]",
            "current_wrong_output_occurrence_1": "заявление Джумабаеву Розу Багиткалиевну",
            "expected_output_occurrence_0": "Джумабаеву Розу Багиткалиевну",
            "expected_case_occurrence_0": "accusative",
            "expected_output_occurrence_1": "заявление Джумабаевой Розы Багиткалиевны",
            "expected_case_occurrence_1": "genitive",
            "reason": "Each occurrence must be analyzed independently based on its own local grammar (verb 'принять [кого]' = accusative vs noun 'заявление [кого]' = genitive), not inherit the case from a previous occurrence of the same placeholder.",
        },
        {
            "id": "regression_fixture_minimal",
            "description": "Minimal fixture combining all cases above for automated regression testing, including capitalization rule for the signature-table job title.",
            "document_structure": [
                {"type": "body_paragraph", "index": 0, "text": "Принять [ФИО] на должность главного менеджера..."},
                {"type": "body_paragraph", "index": 1, "text": "Основание: ..., заявление [ФИО]"},
                {"type": "table_cell", "table_index": 0, "row": 0, "col": 0, "text": "[Должность2]"},
                {"type": "table_cell", "table_index": 0, "row": 0, "col": 1, "text": "[ФИО2]"},
            ],
            "placeholder_values": {
                "ФИО": "Джумабаева Роза Багиткалиевна",
                "Должность2": "члена Правления - заместителя председателя Правления",
                "ФИО2": "Н. Джамышев",
            },
            "expected_occurrences_response": {
                "occurrences": [
                    {"placeholder": "ФИО", "occurrence_index": 0, "source_type": "body_paragraph", "original_value": "Джумабаева Роза Багиткалиевна", "corrected_value": "Джумабаеву Розу Багиткалиевну", "changed": True},
                    {"placeholder": "ФИО", "occurrence_index": 1, "source_type": "body_paragraph", "original_value": "Джумабаева Роза Багиткалиевна", "corrected_value": "Джумабаевой Розы Багиткалиевны", "changed": True},
                    {"placeholder": "Должность2", "occurrence_index": 0, "source_type": "table_cell", "original_value": "члена Правления - заместителя председателя Правления", "corrected_value": "Члена правления - заместителя председателя правления", "changed": True},
                    {"placeholder": "ФИО2", "occurrence_index": 0, "source_type": "table_cell", "original_value": "Н. Джамышев", "corrected_value": "Н. Джамышев", "changed": False},
                ]
            },
            "success_criteria": [
                "Occurrence 0 of ФИО is accusative.",
                "Occurrence 1 of ФИО is genitive, independently derived, not copied from occurrence 0.",
                "Должность2 in the table cell has only its first letter capitalized; 'правления' (second occurrence within the phrase) is lowercase.",
                "ФИО2 in the table cell remains nominative/unchanged regardless of the case used for ФИО elsewhere in the document.",
                "Должность2 in the table cell is not merged with or influenced by the preceding body paragraphs.",
            ],
        },
    ]
}


AI_PLACEHOLDER_ADDITIONAL_EXAMPLES = {
    "examples": [
        {
            "id": "header_invoice_number_no_correction",
            "description": "Document header field (РегНомерДокумента) is a code/number, not natural language — must pass through unchanged regardless of UseAI, and must not be merged with the date field next to it in the same header row.",
            "document_structure": [
                {"type": "header_field", "label": "ДатаДокумента", "text": "<ДатаДокумента> года"},
                {"type": "header_field", "label": "РегНомерДокумента", "text": "№ <РегНомерДокумента>"},
            ],
            "template_placeholders": {
                "ДатаДокумента": "17.06.2026",
                "РегНомерДокумента": "125-ЛС",
            },
            "current_wrong_output": "№ 125 ЛС года",
            "expected_output": {
                "ДатаДокумента": "17 июня 2026",
                "РегНомерДокумента": "125-ЛС",
            },
            "reason": "РегНомерДокумента is an order/registration number (alphanumeric code with a hyphen) — it must never be reworded, expanded into words, or have its punctuation altered, and it must not absorb the word 'года' from the adjacent date field. ДатаДокумента, by contrast, is a real date and may be converted to word form if the surrounding template uses 'года' as a literal trailing word expecting a day+month+year phrase.",
            "success_criteria": [
                "РегНомерДокумента stays exactly '125-ЛС', unchanged.",
                "ДатаДокумента and РегНомерДокумента are corrected independently — no cross-bleed between the two header fields.",
                "No extra words ('года', 'номер', etc.) are injected into РегНомерДокумента.",
            ],
        },
        {
            "id": "city_name_no_declension",
            "description": "Geographic/city name placeholder appearing twice in parallel header columns (Kazakh and Russian) — must not be declined or translated, must remain consistent in both columns.",
            "document_structure": [
                {"type": "header_field", "column": "kazakh", "text": "Алматы қаласы"},
                {"type": "header_field", "column": "russian", "text": "город Алматы"},
            ],
            "template_placeholders": {"Город": "Алматы"},
            "current_wrong_output": None,
            "expected_output": {
                "kazakh_column": "Алматы қаласы",
                "russian_column": "город Алматы",
            },
            "reason": "City names are proper nouns and must stay in nominative case in this kind of bilingual document header regardless of any surrounding case-government words ('қаласы'/'город' are themselves invariant labels meaning 'city of'). The AI must recognize 'Город' here is a header label, not a grammatical object inside a sentence, and must not attempt to decline 'Алматы' (which is indeclinable in Russian anyway) or alter the Kazakh column independently from the Russian one.",
            "success_criteria": [
                "City name unchanged in both language columns.",
                "No declension attempted on an indeclinable proper noun.",
                "Kazakh and Russian columns are not cross-corrected based on each other's grammar.",
            ],
        },
        {
            "id": "organization_name_with_quotes_preserved",
            "description": "Organization name with internal quotation marks and an abbreviation in Latin script (АО «Halyk Finance») embedded inside a body sentence that also needs case correction on a nearby placeholder — verify the org name itself is untouched while the surrounding sentence grammar is still corrected.",
            "context_snippet": "Принять [ФИО] на должность главного менеджера управления учета брокерской деятельности АО «Halyk Finance», на условиях заключенного трудового договора, с <ДатаПриема> года.",
            "template_placeholders": {
                "ФИО": "Джумабаева Роза Багиткалиевна",
                "ДатаПриема": "01.07.2026",
            },
            "current_wrong_output": "Принять Джумабаеву Розу Багиткалиевну на должность главного менеджера управления учета брокерской деятельности АО «Halyk финанс», на условиях заключенного трудового договора, с 1 июля 2026 года.",
            "expected_output": "Принять Джумабаеву Розу Багиткалиевну на должность главного менеджера управления учета брокерской деятельности АО «Halyk Finance», на условиях заключенного трудового договора, с 01 июля 2026 года.",
            "reason": "АО «Halyk Finance» is a legal entity name and must be preserved EXACTLY as written, including Latin script, capitalization, and quotation marks — it is not part of the placeholder set and must never be transliterated, translated, or 'corrected' for spelling. The placeholder ФИО is correctly declined to accusative ('Джумабаеву Розу Багиткалиевну') because 'принять [кого] на должность' governs accusative case. ДатаПриема may be converted to word form for the day+month but should preserve the original numeric day format with leading zero if that is how dates are styled elsewhere in this document (verify against the document's own date convention rather than always stripping leading zeros).",
            "success_criteria": [
                "АО «Halyk Finance» is byte-for-byte identical to the source text — zero edits, since it isn't a placeholder at all.",
                "ФИО occurrence here is accusative, consistent with the verb 'принять ... на должность'.",
                "Correction of one placeholder does not trigger unwanted edits to fixed, non-placeholder text elsewhere in the same sentence.",
            ],
        },
        {
            "id": "contract_basis_clause_multiple_placeholders_one_sentence",
            "description": "A single sentence contains three different placeholders (НомерДоговора, ДатаНачалаДоговора, and a name in genitive) that must each be corrected according to their own grammatical role, without one correction affecting another.",
            "context_snippet": "Основание: трудовой договор № [НомерДоговора] от [ДатаНачалаДоговора] года, заявление [ФИО]",
            "template_placeholders": {
                "НомерДоговора": "45/2026",
                "ДатаНачалаДоговора": "01.07.2026",
                "ФИО": "Джумабаева Роза Багиткалиевна",
            },
            "current_wrong_output": "Основание: трудовой договор № сорок пять от первого июля две тысячи двадцать шестого года, заявление Джумабаеву Розу Багиткалиевну",
            "expected_output": "Основание: трудовой договор № 45/2026 от 01 июля 2026 года, заявление Джумабаевой Розы Багиткалиевны",
            "reason": "НомерДоговора is a contract number/code and must stay in its original numeric/alphanumeric form ('45/2026') — it must NEVER be spelled out in words, even though 'ДатаНачалаДоговора' immediately next to it legitimately gets converted into a word-form date per this document's convention ('01 июля 2026 года'). ФИО here follows 'заявление [кого]' (genitive, 'application OF someone'), which is a DIFFERENT case than the same name's occurrence elsewhere in the document under 'принять [кого] на должность' (accusative) — each occurrence is graded independently by its own local governing word, never by global consistency with other occurrences of the same placeholder.",
            "success_criteria": [
                "НомерДоговора remains a raw alphanumeric code, never spelled out as words.",
                "ДатаНачалаДоговора is converted to word form consistent with the document's date style, independent of how НомерДоговора is (not) converted.",
                "ФИО is genitive here, NOT the same case used for this name's other occurrence in the document.",
                "All three placeholders in the same sentence are corrected independently without one rule incorrectly applying to a different placeholder just because it sits next to a date that does get word-converted.",
            ],
        },
        {
            "id": "kazakh_russian_bilingual_heading_independent_correction",
            "description": "Document title appears in both Kazakh ('БҰЙРЫҚ') and Russian ('ПРИКАЗ') as parallel headings — these are not translations to be cross-checked against each other, and neither should be 'corrected' since they are fixed document-type labels, not placeholders.",
            "document_structure": [
                {"type": "heading", "column": "kazakh", "text": "БҰЙРЫҚ"},
                {"type": "heading", "column": "russian", "text": "ПРИКАЗ"},
            ],
            "template_placeholders": {},
            "current_wrong_output": "БҦЙРЫК / ЗАКАЗ",
            "expected_output": {
                "kazakh_column": "БҰЙРЫҚ",
                "russian_column": "ПРИКАЗ",
            },
            "reason": "Neither heading is a placeholder — both are fixed boilerplate document-type titles in their respective languages. The AI correction step must recognize text with no placeholder markers ([...] or <...>) as out-of-scope entirely and must never 'translate-check' one language column against the other. Also must not alter Kazakh-specific Cyrillic characters (Ұ, Қ) by normalizing them to visually similar but incorrect Russian Cyrillic letters.",
            "success_criteria": [
                "Both headings are completely unmodified — they are not placeholders and should never be sent through AI correction at all.",
                "Kazakh-specific letters (Ұ, Қ, etc.) are preserved exactly, never substituted with similar-looking Russian letters.",
                "No cross-language 'consistency' correction is attempted between the two parallel heading columns.",
            ],
        },
    ]
}

AI_PLACEHOLDER_ALL_EXAMPLES = {
    "examples": AI_PLACEHOLDER_EXAMPLES["examples"] + AI_PLACEHOLDER_ADDITIONAL_EXAMPLES["examples"]
}



AI_PLACEHOLDER_SYSTEM_PROMPT = """Ты — редактор официальных деловых документов на русском языке (приказы, заявления, служебные записки, кадровые документы и т.п.). Твоя задача — исправить значения плейсхолдеров так, чтобы они грамматически и стилистически правильно вписывались в окружающий текст документа, сохраняя при этом исходный смысл и фактическую информацию.

ВХОДНЫЕ ДАННЫЕ:
Ты получишь:
1. Полный текст документа (с плейсхолдерами в виде [ИмяПлейсхолдера]).
2. Список плейсхолдеров с их текущими (возможно неправильными) значениями.
3. Дополнительную инструкцию от пользователя (если есть) — учитывай её как приоритетное указание сверх общих правил ниже.

ПРАВИЛА ИСПРАВЛЕНИЯ:

1. ГРАММАТИКА И ПАДЕЖИ
   Каждое значение должно стоять в правильном падеже согласно своей роли в предложении на месте конкретного вхождения плейсхолдера. Один и тот же плейсхолдер может встречаться несколько раз в документе в разных падежах — анализируй КАЖДОЕ вхождение отдельно по контексту вокруг него, а не один раз для всего ключа.

   Пример:
   "От [ФИОСотрудника]" → родительный падеж → "От Иванова Ивана Ивановича"
   "Прошу предоставить [ФИОСотрудника] отпуск" → дательный падеж → "Прошу предоставить Иванову Ивану Ивановичу отпуск"

2. ОФИЦИАЛЬНО-ДЕЛОВОЙ СТИЛЬ
   Исправляй разговорные, сокращённые или неформальные формулировки на официально-деловые, принятые в кадровом и юридическом делопроизводстве РК/РФ. Например:
   "по семейный обстоятельства" → "по семейным обстоятельствам"
   "директору департамента" (если контекст требует другого падежа/склонения) → "директора департамента по управлению персоналом" или соответствующая правильная форма

3. ДАТЫ
   Приводи даты к официальному формату документа, если контекст требует словесной формы:
   "21.06.2026" → "21 июня 2026 года" (если в окружающем тексте дата упоминается словесно, например "Дата ____ 2026 года")
   Если в документе дата используется в числовом формате (таблица, поле "Дата"), сохраняй числовой формат "дд.мм.гггг" без изменений, ЕСЛИ контекст явно не указывает на словесную форму.

4. ДОЛЖНОСТИ И ОРГАНИЗАЦИОННЫЕ НАЗВАНИЯ
   Сверяй должности/подразделения с правильным склонением и официальным наименованием, не сокращай и не меняй смысл (например, "директор" ≠ "директору", если контекст требует именительного падежа в шапке документа — "Кому: Директору департамента...").

5. ИМЕНА СОБСТВЕННЫЕ
   ФИО, названия организаций, БИН/ИИН, номера документов — НЕ изменяй по существу, только склоняй ФИО под правильный падеж согласно правилу 1. Никогда не выдумывай и не дополняй фактическую информацию, которой нет во входных данных.

6. ЧТО НЕ ТРОГАТЬ
   Если значение уже грамматически и стилистически корректно в данном контексте — верни его БЕЗ ИЗМЕНЕНИЙ. Не переписывай то, что не нуждается в исправлении.

7. ОГРАНИЧЕНИЯ
   - Не добавляй новую фактическую информацию, которой не было в исходном значении.
   - Не меняй структуру документа, не предлагай альтернативные формулировки самого документа — только значения плейсхолдеров.
   - Сохраняй регистр первой буквы исходного значения, если контекст (начало предложения / середина предложения) того требует.

ФОРМАТ ОТВЕТА:
Верни СТРОГО JSON без markdown-разметки, без комментариев, без пояснений — только JSON в следующей структуре:

{
  "occurrences": [
    {
      "placeholder": "ИмяПлейсхолдера",
      "occurrence_index": 0,
      "original_value": "исходное значение",
      "corrected_value": "исправленное значение",
      "changed": true
    }
  ]
}

Где:
- "placeholder" — имя плейсхолдера без квадратных скобок.
- "occurrence_index" — порядковый номер вхождения этого плейсхолдера в документе (0 для первого вхождения, 1 для второго и т.д.), если плейсхолдер встречается несколько раз.
- "original_value" — значение, которое было передано на вход для этого плейсхолдера.
- "corrected_value" — исправленное значение для данного конкретного вхождения.
- "changed" — true, если значение было изменено; false, если оставлено как есть.

Если плейсхолдер встречается только один раз в документе — верни один объект с occurrence_index: 0.

Не возвращай ничего, кроме этого JSON-объекта.

ПРИМЕРЫ ДЛЯ ОРИЕНТАЦИИ:
""" + json.dumps(AI_PLACEHOLDER_ALL_EXAMPLES, ensure_ascii=False, indent=2)


def _openai_placeholder_payload(
    slot_values: dict[str, Any],
    contexts: dict[str, list[str]],
    prompt_ai: str,
    occurrences: list[dict[str, Any]] | None = None,
    full_document_text: str = "",
) -> dict[str, Any]:
    occurrence_list = [item for item in occurrences or [] if not item.get("ai_excluded")]
    skipped = len(occurrences or []) - len(occurrence_list)
    if skipped:
        logger.warning(
            "UseAI OpenAI payload omits %s extracted placeholder occurrences because they are ai_excluded: %s",
            skipped,
            [
                {
                    "placeholder": item.get("placeholder", item.get("key")),
                    "occurrence_index": item.get("occurrence_index"),
                    "source_type": item.get("source_type"),
                    "source_path": item.get("source_path"),
                    "reason": item.get("ai_exclusion_reason"),
                    "context_text": item.get("context_text"),
                }
                for item in occurrences or []
                if item.get("ai_excluded")
            ],
        )
    placeholder_payload = {
        "values": slot_values,
        "occurrences": [
            {
                "placeholder": item.get("placeholder", item.get("key")),
                "occurrence_index": item.get("occurrence_index", 0),
                "original_value": item.get("original_value", item.get("value", "")),
                "source_type": item.get("source_type", ""),
                "source_path": item.get("source_path", ""),
                "context": item.get("context", ""),
                "context_text": item.get("context_text", item.get("context", "")),
                "context_with_value": item.get("context_with_value", ""),
            }
            for item in occurrence_list
        ],
    }
    user_prompt = (
        "ДОКУМЕНТ (полный текст с плейсхолдерами):\n"
        "---\n"
        f"{full_document_text}\n"
        "---\n\n"
        "ПЛЕЙСХОЛДЕРЫ И ЗНАЧЕНИЯ:\n"
        f"{json.dumps(placeholder_payload, ensure_ascii=False)}\n\n"
        "ДОПОЛНИТЕЛЬНАЯ ИНСТРУКЦИЯ ОТ ПОЛЬЗОВАТЕЛЯ:\n"
        f"{prompt_ai or ''}\n\n"
        "Исправь значения плейсхолдеров согласно правилам выше, учитывая контекст каждого конкретного "
        "вхождения в тексте документа."
    )
    return {
        "model": os.environ.get("OPENAI_PLACEHOLDER_MODEL", os.environ.get("OPENAI_MODEL", "gpt-4o-mini")),
        "messages": [
            {"role": "system", "content": AI_PLACEHOLDER_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        "response_format": {"type": "json_object"},
        "temperature": 0,
    }


def _parse_openai_chat_content(raw_response: bytes) -> str:
    payload = json.loads(raw_response.decode("utf-8"))
    choices = payload.get("choices")
    if not isinstance(choices, list) or not choices:
        raise ValueError("OpenAI response has no choices")
    message = choices[0].get("message") if isinstance(choices[0], dict) else None
    content = message.get("content") if isinstance(message, dict) else None
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict) and isinstance(item.get("text"), str):
                parts.append(item["text"])
        if parts:
            return "".join(parts)
    raise ValueError("OpenAI response content is empty")


def _case_hint_for_placeholder_occurrence(key: str, context: str) -> str | None:
    lower_key = key.lower()
    lower_context = context.lower()
    bracket_key = f"[{lower_key}]"
    is_person_key = "фио" in lower_key or "сотрудник" in lower_key
    is_date_key = "дата" in lower_key or "date" in lower_key
    is_number_key = "номер" in lower_key or "number" in lower_key or "code" in lower_key

    if is_number_key and re.search(rf"(?:№|номер\s+)\s*{re.escape(bracket_key)}", lower_context):
        return "preserve"
    if is_date_key and re.search(rf"(?:^|\s)от\s+{re.escape(bracket_key)}\s*(?:года|г\.|год)(?:\s|$|[.,;:])", lower_context):
        return "date_ru_no_year_word"
    if re.search(rf"(?:^|\s)от\s+{re.escape(bracket_key)}(?:\s|$|[.,;:])", lower_context):
        return "gent"
    if is_person_key and re.search(rf"(?:^|\s)заявлени[еяю]\s+{re.escape(bracket_key)}(?:\s|$|[.,;:])", lower_context):
        return "gent"
    if is_person_key and re.search(rf"(?:^|\s)принять\s+{re.escape(bracket_key)}(?:\s|$|[.,;:])", lower_context):
        return "accs"
    if re.search(rf"(?:^|\s)дата\s+{re.escape(bracket_key)}(?:\s|$|[.,;:])", lower_context):
        return "preserve"
    if "предоставить" in lower_context and "отпуск" in lower_context:
        if is_person_key:
            return "datv"
        if "должност" in lower_key:
            return "preserve"
    return None


def _request_ai_placeholder_corrections(
    slot_values: dict[str, Any],
    contexts: dict[str, list[str]],
    prompt_ai: str,
    occurrences: list[dict[str, Any]] | None = None,
    full_document_text: str = "",
    log_key: str | None = None,
    call_log: dict[str, Any] | None = None,
) -> dict[str, str]:
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if call_log is not None:
        call_log["key"] = log_key
        call_log["openai_config"] = {
            "model": os.environ.get("OPENAI_PLACEHOLDER_MODEL", os.environ.get("OPENAI_MODEL", "gpt-4o-mini")),
            "base_url": os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1").rstrip("/"),
        }
    if not api_key:
        if call_log is not None:
            call_log["error"] = "OPENAI_API_KEY is not configured"
        logger.warning(
            "UseAI requested for replace-edit, but OPENAI_API_KEY is not configured: use_ai_log_key=%s",
            log_key,
        )
        return {}

    body = json.dumps(
        _openai_placeholder_payload(slot_values, contexts, prompt_ai, occurrences, full_document_text),
        ensure_ascii=False,
    ).encode("utf-8")
    base_url = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1").rstrip("/")
    request_body_text = body.decode("utf-8", errors="replace")
    if call_log is not None:
        call_log["request"] = {
            "url": f"{base_url}/chat/completions",
            "method": "POST",
            "body": json.loads(request_body_text),
        }
    logger.info(
        "UseAI requested for replace-edit; calling OpenAI placeholder correction: use_ai_log_key=%s model=%s base_url=%s placeholders=%s occurrences=%s prompt_present=%s",
        log_key,
        os.environ.get("OPENAI_PLACEHOLDER_MODEL", os.environ.get("OPENAI_MODEL", "gpt-4o-mini")),
        base_url,
        sorted(str(key) for key in slot_values.keys()),
        len(occurrences or []),
        bool(prompt_ai.strip()),
    )
    logger.info(
        "UseAI OpenAI request body: use_ai_log_key=%s body=%s",
        log_key,
        request_body_text,
    )
    req = Request(
        f"{base_url}/chat/completions",
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urlopen(req, timeout=OPENAI_PLACEHOLDER_TIMEOUT_SECONDS) as resp:
            raw_response = resp.read()
            logger.info(
                "OpenAI placeholder correction response received: use_ai_log_key=%s status=%s bytes=%s",
                log_key,
                getattr(resp, "status", None),
                len(raw_response),
            )
            raw_response_text = raw_response.decode("utf-8", errors="replace")
            logger.info(
                "UseAI OpenAI raw response: use_ai_log_key=%s body=%s",
                log_key,
                raw_response_text,
            )
            if call_log is not None:
                try:
                    response_body: Any = json.loads(raw_response_text)
                except json.JSONDecodeError:
                    response_body = raw_response_text
                call_log["response"] = {
                    "status": getattr(resp, "status", None),
                    "bytes": len(raw_response),
                    "body": response_body,
                }
            content = _parse_openai_chat_content(raw_response)
    except HTTPError as exc:
        raw_error = exc.read()
        raw_error_text = raw_error.decode("utf-8", errors="replace")
        if call_log is not None:
            try:
                error_body: Any = json.loads(raw_error_text)
            except json.JSONDecodeError:
                error_body = raw_error_text
            call_log["response"] = {
                "status": exc.code,
                "bytes": len(raw_error),
                "body": error_body,
            }
            call_log["error"] = f"OpenAI HTTP {exc.code}"
        logger.error(
            "UseAI OpenAI error response: use_ai_log_key=%s status=%s body=%s",
            log_key,
            exc.code,
            raw_error_text,
        )
        raise

    parsed = json.loads(content)
    if not isinstance(parsed, dict):
        raise ValueError("OpenAI correction payload is not a JSON object")
    occurrence_lookup = {
        (str(item.get("placeholder", item.get("key"))), int(item.get("occurrence_index", 0))): str(item.get("id"))
        for item in occurrences or []
        if (item.get("placeholder") or item.get("key")) and not item.get("ai_excluded")
    }
    corrections: dict[str, str] = {}
    parsed_occurrences = parsed.get("occurrences")
    if isinstance(parsed_occurrences, list):
        for item in parsed_occurrences:
            if not isinstance(item, dict):
                continue
            placeholder = str(item.get("placeholder") or "").strip()
            try:
                occurrence_index = int(item.get("occurrence_index", 0))
            except (TypeError, ValueError):
                continue
            correction_id = occurrence_lookup.get((placeholder, occurrence_index))
            corrected = item.get("corrected_value")
            if correction_id and corrected is not None and str(corrected).strip():
                corrections[correction_id] = str(corrected)
        return corrections

    allowed = {str(item.get("id")) for item in occurrences or [] if item.get("id") and not item.get("ai_excluded")} or set(slot_values)
    for key, value in parsed.items():
        if key in allowed and value is not None and str(value).strip():
            corrections[str(key)] = str(value)
    return corrections


def _is_non_trivial_ai_change(original: str, corrected: str) -> bool:
    """True when the GPT change is more substantive than date-format rewriting or pure casing."""
    if original == corrected:
        return False
    if original.lower() == corrected.lower():
        return False
    import re as _re
    if _re.match(r"^\d{1,2}[.\-/]\d{2}[.\-/]\d{4}$", original.strip()):
        return False
    return True


def _ai_correct_slot_values(
    doc: Document,
    slot_values: dict[str, str],
    prompt_ai: str,
    log_key: str | None = None,
    call_log: dict[str, Any] | None = None,
) -> tuple[dict[str, str], dict[tuple[str, int], str], str]:
    """Returns (slot_values, occurrence_values, review_summary).

    GPT and Claude run in parallel when Claude is available.
    Claude wins on any per-occurrence disagreement with GPT.
    """
    import concurrent.futures

    if not slot_values:
        return slot_values, {}, ""

    # Pre-extract once so both AI threads share the same input without re-reading the doc
    occurrences = _ai_extract_placeholder_occurrences(doc, slot_values)
    full_text = _ai_document_full_text(doc)

    claude_is_available = _ai_claude_available()

    gpt_occurrence_values: dict[tuple[str, int], str] = {}
    gpt_per_key: dict[str, str] = {}
    claude_corrections: dict[tuple[str, int], str] = {}
    review_summary = ""
    claude_ran = False
    gpt_result = None

    with concurrent.futures.ThreadPoolExecutor(max_workers=2 if claude_is_available else 1) as pool:
        gpt_future = pool.submit(
            _ai_pipeline_correct_slot_values,
            doc, slot_values, prompt_ai,
            log_key=log_key,
            call_log=call_log,
            timeout_seconds=OPENAI_PLACEHOLDER_TIMEOUT_SECONDS,
            _precomputed_occurrences=occurrences,
            _precomputed_full_text=full_text,
        )

        # Claude gets original occurrences with no GPT pre-answer — fully independent
        claude_future = pool.submit(
            _ai_claude_correct_occurrences,
            full_text, occurrences, {},
            prompt_ai, None, log_key, call_log,
        ) if claude_is_available else None

        try:
            gpt_result = gpt_future.result(timeout=OPENAI_PLACEHOLDER_TIMEOUT_SECONDS + 10)
            gpt_occurrence_values = gpt_result.occurrence_values
            for (key, _occ1), val in gpt_occurrence_values.items():
                gpt_per_key[key] = val
        except Exception as exc:
            logger.exception(
                "GPT placeholder correction failed; using original values: use_ai_log_key=%s error=%s",
                log_key, exc,
            )
            if call_log is not None:
                call_log["error"] = f"GPT correction failed: {exc}"
            if claude_future:
                claude_future.cancel()
            return slot_values, {}, ""

        if claude_future:
            try:
                claude_corrections, review_summary = claude_future.result(timeout=30)
                claude_ran = True
            except Exception as exc:
                logger.warning(
                    "Claude correction pass failed; using GPT output: use_ai_log_key=%s error=%s",
                    log_key, exc,
                )
                review_summary = "Проверка Claude недоступна; применены исправления GPT."

    # Merge: Claude wins per occurrence
    final_occurrence_values = {**gpt_occurrence_values, **claude_corrections}
    final_per_key = dict(gpt_per_key)
    for (key, _occ1), val in final_occurrence_values.items():
        final_per_key[key] = val

    # Changes where Claude's answer differs from GPT's
    changes_from_gpt = [
        {
            "placeholder": k,
            "gpt_value": gpt_per_key.get(k, str(slot_values.get(k, ""))),
            "claude_value": v,
            "reason": "Исправлено Claude",
        }
        for k, v in final_per_key.items()
        if v != gpt_per_key.get(k, str(slot_values.get(k, "")))
    ]

    if changes_from_gpt:
        item_status = "pending"
    elif claude_ran:
        item_status = "logged"
    elif any(_is_non_trivial_ai_change(str(slot_values.get(k, "")), v) for k, v in gpt_per_key.items()):
        item_status = "pending"
    else:
        item_status = "logged"

    try:
        claude_result_for_log = {
            "corrected_values": final_per_key,
            "review_summary": {
                "had_issues": bool(changes_from_gpt),
                "changes_from_gpt": changes_from_gpt,
                "note": review_summary or "GPT и Claude проверили все значения.",
            },
        }
        enqueue_background_review_log(
            original_params={"template": full_text, "placeholders": dict(slot_values)},
            gpt_response=gpt_per_key,
            claude_result=claude_result_for_log,
            document=full_text,
            document_name=str((call_log or {}).get("document_name") or (call_log or {}).get("filename") or ""),
            log_key=log_key or "",
            status=item_status,
        )
    except Exception as persist_exc:
        logger.warning("Failed to persist AI correction to review queue: use_ai_log_key=%s error=%s", log_key, persist_exc)

    return (gpt_result.slot_values if gpt_result else slot_values), final_occurrence_values, review_summary



def _para_align(para) -> str:
    align_map = {"CENTER": "center", "RIGHT": "right", "JUSTIFY": "justify"}
    if para.alignment:
        key = str(para.alignment).split(".")[-1]
        return align_map.get(key, "")
    return ""


def _html_escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _table_param_to_html(rows: list[list[str]], key: str) -> str:
    """Render a table-param as a read-only HTML table in the document preview."""
    key_attr = _html_escape(key)
    if not rows:
        return f'<p class="table-param-empty" data-key="{key_attr}">[таблица: {key_attr} — пуста]</p>'
    header, *body_rows = rows
    th_cells = "".join(f"<th>{_html_escape(str(c))}</th>" for c in header)
    tbody = ""
    for row in body_rows:
        td_cells = "".join(f"<td>{_html_escape(str(c))}</td>" for c in row)
        tbody += f"<tr>{td_cells}</tr>"
    return (
        f'<div class="table-param-block" data-key="{key_attr}">'
        f'<table class="doc-table table-param">'
        f"<thead><tr>{th_cells}</tr></thead>"
        f"<tbody>{tbody}</tbody>"
        f"</table></div>"
    )


def _is_sole_placeholder(full_text: str) -> re.Match | None:
    """Return the match if the paragraph contains exactly one placeholder and nothing else."""
    stripped = full_text.strip()
    if not stripped:
        return None
    matches = list(_PLACEHOLDER_RE.finditer(stripped))
    if len(matches) == 1 and matches[0].group(0) == stripped:
        return matches[0]
    return None


def _para_to_html(para, params: dict[str, str], table_params: dict[str, list] | None = None) -> str:
    table_params = table_params or {}
    full_text = _para_full_text(para)
    if not full_text.strip():
        return "<p>&nbsp;</p>"

    # If the whole paragraph is a single table placeholder → render table
    sole = _is_sole_placeholder(full_text)
    if sole:
        key = _match_key(sole)
        if key in table_params:
            return _table_param_to_html(table_params[key], key)

    # Escape HTML chars first ({, } and [, ] are NOT HTML-special so they stay intact)
    escaped = _html_escape(full_text)

    def make_slot(m: re.Match) -> str:
        key = _match_key(m)
        # Dot-notation table cell/row/col reference (e.g. {{Table.1.2}})
        if "." in key:
            resolved = _resolve_table_cell(key, table_params)
            if resolved is not None:
                key_attr = _html_escape(key)
                val_esc = _html_escape(resolved)
                return (
                    f'<span class="param-slot" data-key="{key_attr}" '
                    f'data-original-key="{key_attr}" contenteditable="true" '
                    f'spellcheck="false">{val_esc}</span>'
                )
        if key in table_params:
            # Inline full-table placeholder — show a non-editable badge
            key_attr = _html_escape(key)
            row_count = len(table_params[key])
            col_count = len(table_params[key][0]) if table_params[key] else 0
            return (
                f'<span class="param-slot table-slot" data-key="{key_attr}" '
                f'title="{key_attr}" contenteditable="false">'
                f'📋 {key_attr} ({row_count} стр. × {col_count} кол.)</span>'
            )
        value = params.get(key, "")
        key_attr = _html_escape(key)
        val_esc = _html_escape(value)
        return (
            f'<span class="param-slot" data-key="{key_attr}" '
            f'data-original-key="{key_attr}" contenteditable="true" '
            f'spellcheck="false">{val_esc}</span>'
        )

    inner = _PLACEHOLDER_RE.sub(make_slot, escaped)

    align = _para_align(para)
    style_attr = f' style="text-align:{align}"' if align else ""

    # Bold / italic from first non-empty run
    is_bold = any(r.bold for r in para.runs if r.text.strip())
    is_italic = all(r.italic for r in para.runs if r.text.strip()) and para.runs

    if is_bold:
        inner = f"<strong>{inner}</strong>"
    if is_italic:
        inner = f"<em>{inner}</em>"

    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        parts = style_name.split()
        level = parts[-1] if parts and parts[-1].isdigit() else "2"
        return f"<h{level}{style_attr}>{inner}</h{level}>"

    return f"<p{style_attr}>{inner}</p>"


def docx_to_html(doc: Document, params: dict[str, str], table_params: dict[str, list] | None = None) -> str:
    table_params = table_params or {}
    parts: list[str] = []
    for idx, child in enumerate(doc.element.body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph
            inner = _para_to_html(Paragraph(child, doc), params, table_params)
        elif tag == "tbl":
            from docx.table import Table
            table = Table(child, doc)
            rows_html = []
            for row in table.rows:
                cells_html = []
                for cell in row.cells:
                    cell_content = "".join(
                        _para_to_html(p, params, table_params) for p in cell.paragraphs
                    )
                    cells_html.append(f"<td>{cell_content}</td>")
                rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
            inner = f'<table class="doc-table"><tbody>{"".join(rows_html)}</tbody></table>'
        else:
            continue
        parts.append(f'<div class="doc-body-el" data-body-idx="{idx}">{inner}</div>')
    return "\n".join(parts)


def _insert_word_table_after(doc: Document, ref_el: Any, rows: list[list[str]]) -> None:
    """Insert a new Word table immediately after ref_el in the body."""
    if not rows:
        return
    cols = max((len(r) for r in rows), default=1)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data[:cols]):
            cell = tbl.rows[i].cells[j]
            cell.text = str(cell_text)
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    tbl_el = tbl._element
    tbl_el.getparent().remove(tbl_el)
    ref_el.addnext(tbl_el)


def _insert_word_table(doc: Document, para_el: Any, rows: list[list[str]]) -> None:
    """Replace an XML paragraph element with a Word table built from rows data."""
    if not rows:
        parent = para_el.getparent()
        if parent is not None:
            parent.remove(para_el)
        return
    cols = max((len(r) for r in rows), default=1)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data[:cols]):
            cell = tbl.rows[i].cells[j]
            cell.text = str(cell_text)
            # Bold the header row
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    # Detach the table from where add_table() appended it and move it to the placeholder position
    tbl_el = tbl._element
    tbl_el.getparent().remove(tbl_el)
    para_el.addprevious(tbl_el)
    para_el.getparent().remove(para_el)


_NEW_ROW_MARKER = "НоваяСтрока"


def _insert_table_row_after(table, after_idx: int, template_row_xml):
    row_xml = copy.deepcopy(template_row_xml)
    table.rows[after_idx]._tr.addnext(row_xml)
    return table.rows[after_idx + 1]


def _clear_row_text(row) -> None:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""


def _run_index_at(spans: list[tuple[int, int]], pos: int) -> tuple[int, int] | None:
    for idx, (start, end) in enumerate(spans):
        if start <= pos < end:
            return idx, pos - start
    return None


def _replace_matches_preserving_runs(para, resolver) -> None:
    original_texts = [run.text for run in para.runs]
    full = "".join(original_texts)
    matches = list(_PLACEHOLDER_RE.finditer(full))
    if not matches:
        return

    spans: list[tuple[int, int]] = []
    cursor = 0
    for value in original_texts:
        start = cursor
        cursor += len(value)
        spans.append((start, cursor))

    updated = list(original_texts)
    for match in reversed(matches):
        start_pos = _run_index_at(spans, match.start())
        end_pos = _run_index_at(spans, match.end() - 1)
        if start_pos is None or end_pos is None:
            continue

        start_idx, start_offset = start_pos
        end_idx, end_offset_inclusive = end_pos
        end_offset = end_offset_inclusive + 1
        replacement = str(resolver(match))

        if start_idx == end_idx:
            current = updated[start_idx]
            updated[start_idx] = current[:start_offset] + replacement + current[end_offset:]
            continue

        updated[start_idx] = updated[start_idx][:start_offset] + replacement
        for idx in range(start_idx + 1, end_idx):
            updated[idx] = ""
        updated[end_idx] = updated[end_idx][end_offset:]

    for run, value in zip(para.runs, updated):
        run.text = value


def _replace_in_para(para, resolver) -> None:
    def resolve_match(match: re.Match) -> str:
        return str(resolver(_match_key(match), match.group(0)))

    _replace_matches_preserving_runs(para, resolve_match)


def _replace_in_row(row, resolver) -> None:
    for cell in row.cells:
        for para in cell.paragraphs:
            _replace_in_para(para, resolver)


def _row_text(row) -> str:
    return "\n".join(
        "".join(run.text for run in para.runs)
        for cell in row.cells
        for para in cell.paragraphs
    )


def _normalize_table_object_rows(raw_rows: list) -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    for item in raw_rows:
        if isinstance(item, dict):
            out.append({str(k): str(v) for k, v in item.items()})
    return out


def _expand_object_table_rows(doc: Document, table_object_params: dict[str, list[dict[str, str]]]) -> None:
    if not table_object_params:
        return

    for table in doc.tables:
        row_idx = 0
        while row_idx < len(table.rows):
            row = table.rows[row_idx]
            text = _row_text(row)
            matched_table = None
            for table_name in table_object_params:
                marker = f"{table_name}.{_NEW_ROW_MARKER}."
                if marker in text:
                    matched_table = table_name
                    break

            if not matched_table:
                row_idx += 1
                continue

            template_idx = row_idx
            rows = table_object_params.get(matched_table) or []
            if not rows:
                _clear_row_text(table.rows[template_idx])
                row_idx += 1
                continue

            template_row_xml = copy.deepcopy(table.rows[template_idx]._tr)
            for data_idx, row_data in enumerate(rows):
                target_row = (
                    table.rows[template_idx]
                    if data_idx == 0
                    else _insert_table_row_after(table, template_idx + data_idx - 1, template_row_xml)
                )

                def resolve_row_key(key: str, original: str) -> str:
                    prefix = f"{matched_table}.{_NEW_ROW_MARKER}."
                    if key.startswith(prefix):
                        field_name = key[len(prefix):]
                        return row_data.get(field_name, "")
                    return original

                _replace_in_row(target_row, resolve_row_key)

            row_idx = template_idx + len(rows)


def fill_docx(
    template_bytes: bytes,
    slot_values: dict[str, str],
    table_params: dict[str, list] | None = None,
    injected_tables: list[dict] | None = None,
    table_object_params: dict[str, list[dict[str, str]]] | None = None,
    slot_occurrence_values: dict[tuple[str, int], str] | None = None,
) -> bytes:
    doc = Document(BytesIO(template_bytes))
    table_params = table_params or {}
    injected_tables = injected_tables or []
    table_object_params = table_object_params or {}
    slot_occurrence_values = slot_occurrence_values or {}
    occurrence_counts: dict[str, int] = {}

    _expand_object_table_rows(doc, table_object_params)

    # ── Pass 0: insert user-dragged tables at chosen body positions ────────────
    if injected_tables:
        body_snapshot = list(doc.element.body)
        # Process in reverse order so earlier insertions don't shift later indices
        for inj in sorted(injected_tables, key=lambda x: int(x.get("body_idx", -1)), reverse=True):
            key = str(inj.get("key", ""))
            body_idx = int(inj.get("body_idx", -1))
            if key not in table_params or body_idx < 0 or body_idx >= len(body_snapshot):
                continue
            _insert_word_table_after(doc, body_snapshot[body_idx], table_params[key])

    # ── Pass 1: collect table-placeholder paragraphs (snapshot before mutation) ──
    tbl_targets: list[tuple[Any, str]] = []  # (para_element, key)

    def _collect(paragraphs) -> None:
        for para in paragraphs:
            full = "".join(r.text for r in para.runs)
            sole = _is_sole_placeholder(full)
            if sole:
                key = _match_key(sole)
                if key in table_params:
                    tbl_targets.append((para._element, key))

    _collect(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _collect(cell.paragraphs)

    # ── Pass 2: replace table placeholders with real Word tables ──────────────
    # Process in reverse document order so earlier insertions don't shift positions
    for para_el, key in reversed(tbl_targets):
        _insert_word_table(doc, para_el, table_params[key])

    # ── Pass 3: normal string replacement on all remaining paragraphs ─────────
    def _replace_para(para) -> None:
        full = "".join(r.text for r in para.runs)
        if not _PLACEHOLDER_RE.search(full):
            return

        def _resolve(m: re.Match) -> str:
            key = _match_key(m)
            occurrence_counts[key] = occurrence_counts.get(key, 0) + 1
            occurrence_value = slot_occurrence_values.get((key, occurrence_counts[key]))
            if occurrence_value is not None:
                return occurrence_value
            # User-edited slot values take highest priority
            if key in slot_values:
                return slot_values[key]
            # Dot-notation table cell/row/col reference
            if "." in key:
                resolved = _resolve_table_cell(key, table_params)
                if resolved is not None:
                    return resolved
            return m.group(0)  # leave unknown placeholders as-is

        _replace_matches_preserving_runs(para, _resolve)

    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            _replace_para(Paragraph(child, doc))
        elif tag == "tbl":
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_para(para)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@word_constructor.get("/admin")
def admin_index():
    if _admin_logged_in():
        return redirect("/services/word-constructor/admin/cabinet")
    return redirect("/services/word-constructor/admin/login")


@word_constructor.get("/admin/login")
def admin_login_page():
    return render_template("word_constructor/admin_login.html", error="")


@word_constructor.post("/admin/login")
def admin_login():
    username = request.form.get("username", "")
    password = request.form.get("password", "")
    if not _admin_credentials_ok(username, password):
        return render_template("word_constructor/admin_login.html", error="Invalid username or password"), 401
    session["admin_logged_in"] = True
    return redirect("/services/word-constructor/admin/cabinet")


@word_constructor.post("/admin/logout")
def admin_logout():
    session.pop("admin_logged_in", None)
    return redirect("/services/word-constructor/admin/login")


@word_constructor.get("/admin/cabinet")
def admin_cabinet():
    if not _admin_logged_in():
        return redirect("/services/word-constructor/admin/login")
    with _CLIENT_STORE_LOCK:
        store = _read_client_store()
        clients = [_client_public(client) for client in store["clients"]]
    clients.sort(key=lambda c: c.get("created_at") or "", reverse=True)
    new_token = session.pop("new_client_token", None)
    password_message = session.pop("password_message", "")
    password_error = session.pop("password_error", "")
    try:
        pending_review_count = sum(1 for item in _load_review_items() if item.get("status") == "pending")
    except Exception:
        pending_review_count = 0
    return render_template(
        "word_constructor/admin_cabinet.html",
        clients=clients,
        new_token=new_token,
        admin_username=os.environ.get("ADMIN_USERNAME", "admin"),
        password_message=password_message,
        password_error=password_error,
        pending_review_count=pending_review_count,
    )


@word_constructor.post("/admin/password")
def admin_change_password():
    if not _admin_logged_in():
        return redirect("/services/word-constructor/admin/login")

    current_password = request.form.get("current_password", "")
    new_password = request.form.get("new_password", "")
    confirm_password = request.form.get("confirm_password", "")
    username = os.environ.get("ADMIN_USERNAME", "admin")

    if not _admin_credentials_ok(username, current_password):
        session["password_error"] = "Current password is incorrect"
    elif len(new_password) < 8:
        session["password_error"] = "New password must be at least 8 characters"
    elif new_password != confirm_password:
        session["password_error"] = "New passwords do not match"
    else:
        _set_admin_password(new_password)
        session["password_message"] = "Admin password changed"

    return redirect("/services/word-constructor/admin/cabinet")


@word_constructor.post("/admin/clients")
def admin_create_client():
    if not _admin_logged_in():
        return redirect("/services/word-constructor/admin/login")
    name = (request.form.get("name") or "").strip()
    if not name:
        name = "Service client"
    try:
        expires_at = _parse_admin_expires(request.form.get("expires_at", ""))
    except ValueError:
        expires_at = None

    token = "wc_" + secrets.token_urlsafe(32)
    client = {
        "id": str(uuid.uuid4()),
        "name": name,
        "token_hash": _hash_token(token),
        "created_at": _utc_now_iso(),
        "expires_at": expires_at,
        "active": True,
        "stats": {
            "calls": 0,
            "input_bytes": 0,
            "output_bytes": 0,
            "last_call_at": None,
            "last_path": None,
        },
    }
    with _CLIENT_STORE_LOCK:
        store = _read_client_store()
        store["clients"].append(client)
        _write_client_store(store)
    session["new_client_token"] = token
    return redirect("/services/word-constructor/admin/cabinet")


@word_constructor.post("/admin/clients/<client_id>/toggle")
def admin_toggle_client(client_id: str):
    if not _admin_logged_in():
        return redirect("/services/word-constructor/admin/login")
    with _CLIENT_STORE_LOCK:
        store = _read_client_store()
        for client in store["clients"]:
            if client.get("id") == client_id:
                client["active"] = not bool(client.get("active", True))
                break
        _write_client_store(store)
    return redirect("/services/word-constructor/admin/cabinet")


@word_constructor.get("/")
def index():
    if request_has_client_token():
        return client_api_index_response()
    if not _admin_logged_in():
        return redirect("/services/word-constructor/admin/login")
    return render_template("word_constructor/index.html")


def _docx_to_lines(template_bytes: bytes) -> list[str]:
    """Extract plain-text lines from a .docx for pre-loading in the template builder."""
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.table import Table as DocxTable
    doc = Document(BytesIO(template_bytes))
    lines: list[str] = []
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            lines.append(DocxParagraph(child, doc).text)
        elif tag == "tbl":
            tbl = DocxTable(child, doc)
            for row in tbl.rows:
                cells = [cell.text.replace("\n", " ") for cell in row.cells]
                lines.append("\t".join(cells))
    return lines


def _build_template_docx(lines: list[str] | None = None) -> bytes:
    doc = Document()
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    normalized_lines = [str(line) for line in (lines or [])]
    if normalized_lines:
        for line in normalized_lines:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@word_constructor.get("/template-builder")
def template_builder_page():
    """Visual template builder backed by a blank OnlyOffice session."""
    session_id = str(uuid.uuid4())
    sdir = _session_dir(session_id)
    sdir.mkdir(parents=True, exist_ok=True)

    expires_at = time.time() + SESSION_TB_TTL_SECONDS
    meta = {
        "id": session_id,
        "type": "template_builder",
        "params": [],
        "filename": "template.docx",
        "last_saved_at": time.time(),
        "expires_at": expires_at,
        "expires_at_iso": datetime.fromtimestamp(expires_at, timezone.utc).isoformat(),
    }
    _write_meta(session_id, meta)
    _session_template_path(session_id).write_bytes(_build_template_docx())
    query = request.query_string.decode("utf-8")
    target = f"/services/word-constructor/template-builder/{session_id}"
    if query:
        target = f"{target}?{query}"
    return redirect(target)


@word_constructor.post("/api/template-builder/create")
def api_template_builder_create():
    """
    1C endpoint — create a template-builder session.
    multipart/form-data:
      params   — JSON array of param names: ["ФИО", "Должность", ...]
      document — (optional) base .docx whose text is pre-loaded into the editor
    Returns: {"id": "...", "builder_url": "...", "expires_at": "..."}
    """
    if request.is_json:
        payload = request.get_json(silent=True) or {}
        parsed = payload.get("params", [])
    else:
        params_raw = request.form.get("params", "[]").strip()
        try:
            parsed = json.loads(params_raw)
        except json.JSONDecodeError as exc:
            return jsonify({"error": f"'params' is not valid JSON: {exc}",
                            "example": '["ФИО","Должность","Дата"]'}), 400

    if isinstance(parsed, dict):
        params_list = list(parsed.keys())          # accept object → use keys
    elif isinstance(parsed, list):
        params_list = [str(p) for p in parsed if p]
    else:
        return jsonify({"error": "'params' must be a JSON array or object"}), 400

    # Parse optional base document
    lines: list[str] = []
    template_bytes = None  # type: Optional[bytes]
    filename = "template.docx"
    if request.is_json:
        payload = request.get_json(silent=True) or {}
        filename = str(payload.get("filename", filename) or filename)
        content_base64 = str(payload.get("content_base64", "") or "").strip()
        if content_base64:
            try:
                doc_bytes = _safe_b64decode(content_base64)
            except Exception:
                return jsonify({"error": "Invalid content_base64"}), 400
            try:
                lines = _docx_to_lines(doc_bytes)
                template_bytes = doc_bytes
            except Exception as exc:
                return jsonify({"error": f"Cannot read 'content_base64' as .docx: {exc}"}), 400
    else:
        doc_file = request.files.get("document")
        if doc_file:
            filename = doc_file.filename or filename
            doc_bytes = doc_file.read()
            if doc_bytes:
                try:
                    lines = _docx_to_lines(doc_bytes)
                    template_bytes = doc_bytes
                except Exception as exc:
                    return jsonify({"error": f"Cannot read 'document' as .docx: {exc}"}), 400

    session_id = str(uuid.uuid4())
    sdir = _session_dir(session_id)
    sdir.mkdir(parents=True, exist_ok=True)

    expires_at = time.time() + SESSION_TB_TTL_SECONDS
    meta = {
        "id": session_id,
        "type": "template_builder",
        "params": params_list,
        "filename": filename,
        "last_saved_at": time.time(),
        "expires_at": expires_at,
        "expires_at_iso": datetime.fromtimestamp(expires_at, timezone.utc).isoformat(),
    }
    _write_meta(session_id, meta)
    _session_template_path(session_id).write_bytes(template_bytes or _build_template_docx(lines))

    return jsonify({
        "id": session_id,
        "builder_url": f"/services/word-constructor/template-builder/{session_id}",
        "expires_at": meta["expires_at_iso"],
    })


_DELETED_PAGE = """<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>KazUni — Файл удалён</title>
  <style>
    *{{box-sizing:border-box;margin:0;padding:0}}
    body{{font-family:"Segoe UI",Tahoma,Arial,sans-serif;background:#eef2f6;
          display:flex;align-items:center;justify-content:center;min-height:100vh}}
    .card{{background:#fff;border-radius:12px;box-shadow:0 10px 30px rgba(15,23,42,.08);
           padding:48px 56px;max-width:480px;width:100%;text-align:center}}
    .icon{{font-size:56px;margin-bottom:20px}}
    h1{{font-size:22px;color:#142433;margin-bottom:12px}}
    p{{color:#5c6b7a;line-height:1.6;font-size:15px}}
  </style>
</head>
<body>
  <div class="card">
    <div class="icon">✅</div>
    <h1>Файл успешно загружен</h1>
    <p>Шаблон был скачан системой 1С и сессия завершена.<br>
       Для продолжения работы отправьте новый запрос из 1С.</p>
  </div>
</body>
</html>"""


@word_constructor.get("/template-builder/<session_id>")
def template_builder_session(session_id: str):
    """Template builder backed by an OnlyOffice editing session."""
    meta = _read_meta(session_id)
    if meta is None:
        from flask import make_response
        return make_response(_DELETED_PAGE, 200)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        from flask import make_response
        return make_response(_DELETED_PAGE, 200)
    if meta.get("type") != "template_builder":
        abort(404)
    base = "/services/word-constructor"
    ws_url = f"{public_ws_base_url(request)}{base}/api/template-builder/{session_id}/ws"
    return render_template(
        "word_constructor/template_builder.html",
        session_id=session_id,
        session_params=meta.get("params", []),
        filename=meta.get("filename", "template.docx"),
        expires_at=meta["expires_at_iso"],
        builder_html=meta.get("builder_html", ""),
        websocket_url=ws_url,
        onlyoffice_api_url=_onlyoffice_api_url(),
    )


@word_constructor.get("/api/template-builder/<session_id>/onlyoffice/file")
def api_template_builder_onlyoffice_file(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        abort(410)

    path = _session_template_path(session_id)
    if not path.exists():
        abort(404)

    return _send_file_compat(
        path,
        as_attachment=False,
        download_name=meta.get("filename", "template.docx"),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@word_constructor.get("/api/template-builder/<session_id>/onlyoffice/config")
def api_template_builder_onlyoffice_config(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        abort(410)

    path = _session_template_path(session_id)
    if not path.exists():
        abort(404)

    stat = path.stat()
    filename = meta.get("filename", "template.docx")
    extension = path.suffix.lower().lstrip(".") or "docx"
    file_url = _builder_internal_url(session_id, "onlyoffice/file")
    callback_url = _builder_internal_url(session_id, "onlyoffice/callback")
    editor_key = meta.get("editor_key") or _builder_editor_key(session_id, path)
    meta["editor_key"] = editor_key
    _write_meta(session_id, meta)

    payload = {
        "document": {
            "fileType": extension,
            "key": editor_key,
            "title": filename,
            "url": file_url,
        },
        "documentType": _onlyoffice_document_type(extension),
        "editorConfig": {
            "callbackUrl": callback_url,
            "lang": "ru",
            "customization": {
                "autosave": True,
                "compactHeader": False,
                "forcesave": True,
                "uiTheme": "theme-white",
            },
            "mode": "edit",
            "user": {
                "id": "template-builder",
                "name": "Template Builder",
            },
        },
        "height": "100%",
        "type": "desktop",
        "width": "100%",
    }
    # Inject sidebar plugins
    _params_plugin_url = (
        f"{public_base_url(request)}"
        f"/services/word-constructor/onlyoffice-plugin/config.json?session_id={session_id}"
    )
    _params_plugin_guid = "asc.{b3c7e1a2-4d5f-4890-bcde-f12345678901}"
    payload["editorConfig"]["plugins"] = {
        "pluginsData": [_params_plugin_url],
        "autostart": [_params_plugin_guid],
    }

    payload["token"] = jwt.encode(payload, _onlyoffice_jwt_secret(), algorithm="HS256")
    return jsonify(payload)


@word_constructor.post("/api/template-builder/<session_id>/onlyoffice/callback")
def api_template_builder_onlyoffice_callback(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": 1, "message": "Session expired"}), 410

    payload = request.get_json(silent=True) or {}
    status = payload.get("status")
    if status not in {2, 3, 6, 7}:
        return jsonify({"error": 0})

    if status in {2, 6}:
        download_url = payload.get("url")
        if not download_url:
            return jsonify({"error": 1, "message": "Missing file URL"})

        normalized_url, host_header = _normalize_callback_download_url(download_url)
        content = _download_remote_file(normalized_url, host_header=host_header)
        _session_template_path(session_id).write_bytes(content)
        meta["last_saved_at"] = time.time()
        _write_meta(session_id, meta)

    return jsonify({"error": 0})


@word_constructor.post("/api/template-builder/<session_id>/document")
def api_template_builder_replace_document(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    doc_file = request.files.get("document")
    if not doc_file:
        return jsonify({"error": "Missing 'document' file field"}), 400

    doc_bytes = doc_file.read()
    if not doc_bytes:
        return jsonify({"error": "Uploaded document is empty"}), 400

    try:
        Document(BytesIO(doc_bytes))
    except Exception as exc:
        return jsonify({"error": f"Cannot read 'document' as .docx: {exc}"}), 400

    meta["filename"] = doc_file.filename or meta.get("filename", "template.docx")
    _session_template_path(session_id).write_bytes(doc_bytes)
    meta["editor_key"] = _builder_editor_key(session_id, _session_template_path(session_id))
    meta["last_saved_at"] = time.time()
    _write_meta(session_id, meta)
    return jsonify({"ok": True, "filename": meta["filename"]})


@word_constructor.post("/api/template-builder/<session_id>/params")
def api_template_builder_update_params(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    data = request.get_json(silent=True) or {}
    params = data.get("params")
    if not isinstance(params, list):
        return jsonify({"error": "Expected 'params' array"}), 400

    meta["params"] = [str(item).strip() for item in params if str(item).strip()]
    _write_meta(session_id, meta)
    return jsonify({"ok": True, "params": meta["params"]})


@word_constructor.post("/api/template-builder/<session_id>/content")
def api_template_builder_save_content(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    data = request.get_json(silent=True) or {}
    html = str(data.get("html", "") or "")
    filename = str(data.get("filename", "") or "").strip()
    if not html.strip():
        return jsonify({"error": "Expected non-empty 'html'"}), 400

    meta["builder_html"] = html
    if filename:
        meta["filename"] = filename
    _write_meta(session_id, meta)
    return jsonify({"ok": True})


@word_constructor.post("/api/template-builder/<session_id>/forcesave")
def api_template_builder_forcesave(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    path = _session_template_path(session_id)
    if not path.exists():
        abort(404)

    key = str(meta.get("editor_key") or _builder_editor_key(session_id, path))
    previous_saved_at = meta.get("last_saved_at")
    try:
        result = _builder_forcesave(session_id, key)
    except Exception as exc:
        return jsonify({"error": f"Force save failed: {exc}"}), 502

    error_code = int(result.get("error", 0) or 0)
    if error_code not in {0, 4}:
        return jsonify({"error": "ONLYOFFICE rejected force save", "details": result}), 502

    if error_code == 0:
        saved = _wait_for_builder_save(session_id, previous_saved_at)
        if not saved:
            return jsonify({"error": "Force save timed out waiting for callback"}), 504

    return jsonify({"ok": True, "result": result})


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_NS = "http://www.w3.org/XML/1998/namespace"


def _insert_bookmarks_in_docx(docx_bytes: bytes) -> bytes:
    """
    Scan word/document.xml for {{ParamName}} placeholders and wrap each
    occurrence in a Word bookmark named ParamName.

    Uses lxml to parse/serialize so all original namespace prefixes are
    preserved exactly — ElementTree rewrites them as ns0/ns1 which corrupts
    the file for Word.

    Handles placeholders split across multiple <w:r> runs by concatenating
    run texts per paragraph, finding matches, then rebuilding runs with
    bookmarkStart / bookmarkEnd elements around each placeholder run.

    Returns modified .docx bytes (same ZIP structure, only document.xml changed).
    """
    from lxml import etree as _ET

    # Match both {{ParamName}} and [ParamName] — the plugin inserts [key] format
    PLACEHOLDER = re.compile(r'\{\{([^}\n\r]{1,80})\}\}|\[([^\[\]\n\r]{1,80})\]')
    W  = _W_NS
    XS = _XML_NS
    TAG_R   = f"{{{W}}}r"
    TAG_RPR = f"{{{W}}}rPr"
    TAG_T   = f"{{{W}}}t"
    TAG_BKS = f"{{{W}}}bookmarkStart"
    TAG_BKE = f"{{{W}}}bookmarkEnd"
    PRESERVE = f"{{{XS}}}space"

    def make_run(rpr_elem, text: str):
        r = _ET.Element(TAG_R)
        if rpr_elem is not None:
            r.append(copy.deepcopy(rpr_elem))
        t = _ET.SubElement(r, TAG_T)
        t.text = text
        if text and (text[0] == " " or text[-1] == " "):
            t.set(PRESERVE, "preserve")
        return r

    def make_bk_start(bk_id: int, name: str):
        e = _ET.Element(TAG_BKS)
        e.set(f"{{{W}}}id", str(bk_id))
        e.set(f"{{{W}}}name", name)
        return e

    def make_bk_end(bk_id: int):
        e = _ET.Element(TAG_BKE)
        e.set(f"{{{W}}}id", str(bk_id))
        return e

    def process_para(para, bk_id: int) -> int:
        children = list(para)
        run_info = [(i, ch) for i, ch in enumerate(children) if ch.tag == TAG_R]
        if not run_info:
            return bk_id

        run_texts: list[str] = []
        run_starts: list[int] = []
        pos = 0
        for _, run in run_info:
            t = run.find(TAG_T)
            text = (t.text or "") if t is not None else ""
            run_texts.append(text)
            run_starts.append(pos)
            pos += len(text)

        concat = "".join(run_texts)
        if "{{" not in concat and "[" not in concat:
            return bk_id

        matches = list(PLACEHOLDER.finditer(concat))
        if not matches:
            return bk_id

        def rpr_at(char_pos: int):
            for k, (_, run) in enumerate(run_info):
                start = run_starts[k]
                end = start + len(run_texts[k])
                if start <= char_pos < end:
                    return run.find(TAG_RPR)
            return run_info[0][1].find(TAG_RPR)

        segments: list[tuple] = []
        last = 0
        for m in matches:
            if m.start() > last:
                segments.append(("text", concat[last:m.start()], last))
            name = m.group(1) or m.group(2)  # group 1 = {{...}}, group 2 = [...]
            segments.append(("ph", concat[m.start():m.end()], name, m.start()))
            last = m.end()
        if last < len(concat):
            segments.append(("text", concat[last:], last))

        new_elems = []
        for seg in segments:
            if seg[0] == "text":
                _, text, char_pos = seg
                if text:
                    new_elems.append(make_run(rpr_at(char_pos), text))
            else:
                _, ph_text, name, char_pos = seg
                new_elems.append(make_bk_start(bk_id, name))
                new_elems.append(make_run(rpr_at(char_pos), ph_text))
                new_elems.append(make_bk_end(bk_id))
                bk_id += 1

        first_run_para_idx = run_info[0][0]
        for _, run in reversed(run_info):
            para.remove(run)
        for k, elem in enumerate(new_elems):
            para.insert(first_run_para_idx + k, elem)

        return bk_id

    # ── Read ZIP ──────────────────────────────────────────────────────────
    with zipfile.ZipFile(BytesIO(docx_bytes)) as zin:
        names = zin.namelist()
        zinfo_map = {info.filename: info for info in zin.infolist()}
        files: dict[str, bytes] = {n: zin.read(n) for n in names}

    if "word/document.xml" not in files:
        return docx_bytes

    # ── Parse with lxml (preserves all namespace prefixes) ───────────────
    root = _ET.fromstring(files["word/document.xml"])

    bk_id = 0
    for para in root.iter(f"{{{W}}}p"):
        bk_id = process_para(para, bk_id)

    modified_xml = _ET.tostring(root, xml_declaration=True,
                                encoding="UTF-8", standalone=True)

    # ── Write back ZIP ────────────────────────────────────────────────────
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            if name == "word/document.xml":
                zout.writestr(zinfo_map[name], modified_xml)
            else:
                zout.writestr(zinfo_map[name], files[name])
    return buf.getvalue()


@word_constructor.get("/api/template-builder/<session_id>/download")
def api_template_builder_download(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        abort(410)

    path = _session_template_path(session_id)
    if not path.exists():
        abort(404)

    filename = meta.get("filename", "template.docx")
    builder_html = str(meta.get("builder_html", "") or "").strip()
    if builder_html:
        doc = _html_to_docx(builder_html)
        buf = BytesIO()
        doc.save(buf)
        file_bytes = buf.getvalue()
    else:
        file_bytes = path.read_bytes()

    # Signal WS handler that download happened, then clean up session
    _tb_ws_notify_downloaded(session_id)
    shutil.rmtree(_session_dir(session_id), ignore_errors=True)

    return _send_file_compat(
        BytesIO(file_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@word_constructor.get("/api/converter/<conversion_id>/source")
def api_converter_source(conversion_id: str):
    meta = _read_meta(conversion_id)
    if meta is None or meta.get("type") != "conversion":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(conversion_id), ignore_errors=True)
        abort(410)

    path = _converter_source_path(conversion_id)
    if not path.exists():
        abort(404)

    return _send_file_compat(
        path,
        as_attachment=False,
        download_name=meta.get("filename", "document.docx"),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@word_constructor.post("/api/1c/converter/word-base64-to-pdf/")
@word_constructor.post("/api/1c/converter/word-base64-to-pdf")
def api_1c_word_base64_to_pdf():
    """
    Convert a base64-encoded .docx payload from 1C to PDF.

    Request JSON:
      {"filename": "template.docx", "content_base64": "..."}

    Response: raw PDF bytes with application/pdf content type.
    """
    payload = request.get_json(silent=True)
    if not isinstance(payload, dict):
        return jsonify({"error": "Expected JSON body"}), 400

    filename = str(payload.get("filename") or "document.docx")
    content_base64 = payload.get("content_base64")
    if not isinstance(content_base64, str) or not content_base64.strip():
        return jsonify({"error": "Missing 'content_base64'"}), 400
    if Path(filename).suffix.lower() != ".docx":
        return jsonify({"error": "Only .docx files can be converted to PDF"}), 400

    try:
        document_bytes = _safe_b64decode(content_base64)
    except Exception as exc:
        return jsonify({"error": f"Invalid base64 document: {exc}"}), 400
    if not document_bytes:
        return jsonify({"error": "Decoded document is empty"}), 400

    try:
        Document(BytesIO(document_bytes))
    except Exception:
        return jsonify({"error": "Cannot read content_base64 as a .docx file"}), 400

    try:
        pdf_bytes = _convert_docx_to_pdf_with_onlyoffice(filename, document_bytes)
    except Exception as exc:
        return jsonify({"error": f"Cannot convert document to PDF: {exc}"}), 502

    pdf_name = f"{Path(filename).stem or 'document'}.pdf"
    return _send_file_compat(
        BytesIO(pdf_bytes),
        as_attachment=True,
        download_name=pdf_name,
        mimetype="application/pdf",
    )


@word_constructor.post("/api/1c/template-builder/bridge")
def api_1c_documents_bridge():
    """
    1C bridge endpoint — create a template-builder session from a JSON body.

    Request body (JSON):
      {
        "filename":       "template.docx",
        "content_base64": "<base64-encoded .docx>",
        "params":         ["ФИО", "Должность", "Дата"]
      }

    Returns:
      {
        "session_id":    "...",
        "builder_url":   "/services/word-constructor/template-builder/<id>?source=from1c",
        "status_url":    "/services/word-constructor/api/template-builder/<id>/status",
        "download_url":  "/services/word-constructor/api/template-builder/<id>/download",
        "expires_at":    "ISO-8601"
      }

    1С пример:
      ТелоЗапроса = Новый Структура;
      ТелоЗапроса.Вставить("filename", "template.docx");
      ТелоЗапроса.Вставить("content_base64", СтрокаBase64);
      ТелоЗапроса.Вставить("params", СписокПараметров);

      ЗаписьJSON = Новый ЗаписьJSON;
      ЗаписьJSON.УстановитьСтроку();
      ЗаписатьJSON(ЗаписьJSON, ТелоЗапроса);
      СтрокаJSON = ЗаписьJSON.Закрыть();

      HTTPЗапрос = Новый HTTPЗапрос("/services/word-constructor/api/1c/template-builder/bridge");
      HTTPЗапрос.УстановитьЗаголовок("Content-Type", "application/json");
      HTTPЗапрос.УстановитьТелоИзСтроки(СтрокаJSON);
      Ответ = HTTPСоединение.ОтправитьДляОбработки(HTTPЗапрос);

      ЧтениеJSON = Новый ЧтениеJSON;
      ЧтениеJSON.УстановитьСтроку(Ответ.ПолучитьТелоКакСтроку());
      Результат = ПрочитатьJSON(ЧтениеJSON);

      УРЛРедактора = Результат.builder_url;
      УРЛСтатуса  = Результат.status_url;
      УРЛСкачать  = Результат.download_url;
    """
    payload = request.get_json(silent=True) or {}

    # ── params ──────────────────────────────────────────────────────────
    raw_params = payload.get("params", [])
    if isinstance(raw_params, dict):
        params_list = list(raw_params.keys())
    elif isinstance(raw_params, list):
        params_list = [str(p) for p in raw_params if p]
    else:
        return jsonify({"error": "'params' must be a JSON array or object"}), 400

    # ── document ─────────────────────────────────────────────────────────
    filename = str(payload.get("filename", "template.docx") or "template.docx")
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    template_bytes = None
    content_base64 = str(payload.get("content_base64", "") or "").strip()
    if content_base64:
        try:
            doc_bytes = _safe_b64decode(content_base64)
        except Exception:
            return jsonify({"error": "Invalid content_base64"}), 400
        try:
            # Validate it's a readable docx
            Document(BytesIO(doc_bytes))
            template_bytes = doc_bytes
        except Exception as exc:
            return jsonify({"error": f"Cannot read content_base64 as .docx: {exc}"}), 400

    # ── create session ────────────────────────────────────────────────────
    session_id = str(uuid.uuid4())
    sdir = _session_dir(session_id)
    sdir.mkdir(parents=True, exist_ok=True)

    expires_at = time.time() + SESSION_TB_TTL_SECONDS
    meta = {
        "id": session_id,
        "type": "template_builder",
        "params": params_list,
        "filename": filename,
        "status": "editing",          # editing | ready
        "last_saved_at": time.time(),
        "expires_at": expires_at,
        "expires_at_iso": datetime.fromtimestamp(expires_at, timezone.utc).isoformat(),
    }
    if getattr(g, "api_client_id", None):
        meta["client_id"] = g.api_client_id
    _write_meta(session_id, meta)
    _session_template_path(session_id).write_bytes(template_bytes or _build_template_docx([]))

    base = "/services/word-constructor"
    return jsonify({
        "session_id":    session_id,
        "builder_url":   f"{base}/template-builder/{session_id}?source=from1c",
        "status_url":    f"{base}/api/template-builder/{session_id}/status",
        "download_url":  f"{base}/api/template-builder/{session_id}/download",
        "complete_url":  f"{base}/api/template-builder/{session_id}/ready",
        "websocket_url": f"{public_ws_base_url(request)}{base}/api/template-builder/{session_id}/ws",
        "expires_at":    meta["expires_at_iso"],
    }), 201


@word_constructor.get("/api/template-builder/<session_id>/status")
def api_template_builder_status(session_id: str):
    """
    Poll this endpoint until status == "ready", then download the template.

    Response:
      { "status": "editing" | "ready" | "expired", "session_id": "..." }

    1С пример опроса:
      Пока Истина Цикл
          Ответ = HTTPСоединение.Получить(Новый HTTPЗапрос("/services/word-constructor/api/template-builder/" + ИдСессии + "/status"));
          ЧтениеJSON = Новый ЧтениеJSON;
          ЧтениеJSON.УстановитьСтроку(Ответ.ПолучитьТелоКакСтроку());
          Данные = ПрочитатьJSON(ЧтениеJSON);
          Если Данные.status = "ready" Тогда
              Прервать;
          КонецЕсли;
      КонецЦикла;
    """
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        return jsonify({"status": "not_found", "session_id": session_id}), 404
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"status": "expired", "session_id": session_id})
    return jsonify({
        "status":     meta.get("status", "editing"),
        "session_id": session_id,
        "expires_at": meta.get("expires_at_iso"),
    })


# In-memory store for WebSocket waiting events per template-builder session
_TB_WS_EVENTS: dict = {}  # session_id → {"event": threading.Event, "payload": dict | None, "download_event": threading.Event}
_TB_WS_LOCK = threading.Lock()
_TB_WS_TIMEOUT = 3600   # seconds to wait for user to submit
_TB_DL_TIMEOUT = 300    # seconds to wait for download after template_ready sent


def _tb_ws_register(session_id: str):
    """Register (or re-register) a WS listener for session_id.

    On reconnect the old entry is reused if a payload is already waiting
    (user clicked «Отправить» while 1C was disconnected), so the new
    connection receives it immediately.  If no payload is ready, a fresh
    event is created so the new wait starts clean.
    """
    import threading as _threading
    with _TB_WS_LOCK:
        existing = _TB_WS_EVENTS.get(session_id)
        if existing and existing.get("payload") is not None:
            # Payload already set — reuse entry so reconnected client gets it
            return existing
        # No payload yet (or no prior entry) — create/reset with fresh events
        _TB_WS_EVENTS[session_id] = {
            "event": _threading.Event(),
            "payload": None,
            "download_event": _threading.Event(),
        }
        return _TB_WS_EVENTS[session_id]


def _tb_ws_notify(session_id: str, payload: dict):
    with _TB_WS_LOCK:
        entry = _TB_WS_EVENTS.get(session_id)
    if entry:
        entry["payload"] = payload
        entry["event"].set()


def _tb_ws_notify_downloaded(session_id: str):
    """Called by the download endpoint to signal the WS handler to close."""
    with _TB_WS_LOCK:
        entry = _TB_WS_EVENTS.get(session_id)
    if entry:
        entry["download_event"].set()
    # Clean up the events entry
    with _TB_WS_LOCK:
        _TB_WS_EVENTS.pop(session_id, None)


@word_constructor.post("/api/template-builder/<session_id>/ready")
def api_template_builder_ready(session_id: str):
    """
    Called by the browser when the user clicks «Отправить в 1С».
    Saves the current HTML content (if provided), marks session status = "ready",
    and notifies any waiting WebSocket listeners.

    Body (JSON, optional):
      { "html": "<p>...</p>", "filename": "my_template.docx" }
    """
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    data = request.get_json(silent=True) or {}
    html = str(data.get("html", "") or "").strip()
    if html:
        meta["builder_html"] = html
    if data.get("filename"):
        meta["filename"] = str(data["filename"])

    meta["status"] = "ready"
    meta["ready_at"] = time.time()
    _write_meta(session_id, meta)

    # Build the docx bytes for WS notification
    base = "/services/word-constructor"
    download_url = f"{base}/api/template-builder/{session_id}/download"

    # Try to produce the file bytes now so WS payload includes content_base64
    file_bytes = None  # type: Optional[bytes]
    try:
        builder_html = str(meta.get("builder_html", "") or "").strip()
        if builder_html:
            doc = _html_to_docx(builder_html)
            buf = BytesIO()
            doc.save(buf)
            file_bytes = buf.getvalue()
        else:
            path = _session_template_path(session_id)
            if path.exists():
                file_bytes = path.read_bytes()
    except Exception:
        pass

    ws_payload: dict = {
        "type":         "template_ready",
        "session_id":   session_id,
        "filename":     meta.get("filename", "template.docx"),
        "download_url": download_url,
    }
    if file_bytes:
        import base64 as _b64
        ws_payload["content_base64"] = _b64.b64encode(file_bytes).decode("ascii")
        ws_payload["size_bytes"] = len(file_bytes)

    _tb_ws_notify(session_id, ws_payload)

    return jsonify({
        "ok":          True,
        "status":      "ready",
        "download_url": download_url,
    })


# ── OnlyOffice sidebar plugin — visual panel for inserting template params ──
# Files are served as routes so session_id can be embedded in the index URL.


@word_constructor.get("/onlyoffice-plugin/config.json")
def oo_plugin_config():
    """Dynamic plugin manifest.

    IMPORTANT: OnlyOffice 9.x always resolves 'url' as relative to config.json —
    absolute URLs get treated as path segments.  We therefore use a bare relative
    path ("index.html") and pass session_id via initData instead of a query param.
    """
    import json as _json
    from flask import Response as _Resp
    session_id = request.args.get("session_id", "")
    cfg = {
        "name": "Параметры 1С",
        "nameLocale": {},
        "guid": "asc.{b3c7e1a2-4d5f-4890-bcde-f12345678901}",
        "version": "2.0.0",
        "minVersion": "6.0.0",
        "variations": [
            {
                "description": "Вставка параметров шаблона KazUni",
                # Relative path only — OO prepends the config.json base dir automatically
                "url": "index.html",
                # session_id delivered via initData so no query-string conflicts
                "initDataType": "text",
                "initData": session_id,
                "icons": ["icon.png"],
                "isViewer": False,
                "EditorsSupport": ["word"],
                "isVisual": True,
                "isModal": False,
                "isInsideMode": True,
                "isUpdateOleOnResize": False,
                "buttons": [],
            }
        ],
    }
    return _Resp(_json.dumps(cfg), mimetype="application/json")


@word_constructor.get("/onlyoffice-plugin/icon.png")
def oo_plugin_icon():
    """1C brand icon (40×40) for the OnlyOffice plugin tab."""
    import base64 as _b64
    from flask import Response as _Resp
    png = _b64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAACgAAAAoCAYAAACM/rhtAAAGI0lEQVR42u2Yb4hcVxnGf++5d2Z2Z2d3srN/wlJhU/xDo1XTBlGwtVZKbUQLhUpbSz/YFtQP2i9KBWlDDNR+UfxXEVEI/eAXi/kYaSVRNGIhQg1qRFqRtEk2uzOzuzM7u3Nn7j2PH+bO7uxmdzbdxCZCXjhw53Lmnud93vc8530P3LD/M5Ow6xVYIBGkz+56A7bKmug8X3OQOohbB6zOp1TjiNr8Q7M8cs1ASrjehVXhPi3xihpI7XRUmVWTWyTCdL69U8DC1d/LTKvMU4qRIqQ5vMrEmiPRPFKDFZU50d04/zOQAusmP4DmuU11jmmRusokqqTAKmh1lPGqICVIZb7WL9x2pRsA8GZIVT5Mjq/T4gFGGGJxzYMtLCGD0WaJEu8yoy5hZuv/Ee40nICZkUhkVeFZMjxFngIRUCXBcKn7tgX1RoCjzVvA8lZr2Q5ENjAjBlCZJzDuZoRHaQBtYowAu8zvioRdBCxyiBLf7jq9IwZ76I81x36GuAfH8wAskgAO9zYiIkSIEfFPoA4EsB7cZQHsAjNDqjOJeBw4SJYBFokRhq1tks120RYx8wwSUOPXNsF3JQLs0ozdVijNkE6R0Tz7aPAkw3zHN8kxTwKEiGDVb23gQEAGyKZU5HoAG45lEhx36hR5MxI2kRvXjzkJ0wy7uZnHqed3M8EBoO0G8QQpsAJQShceAMZX+YEMJGVH/SdD1H5UIP5P0AvS8DgcH+MDHNMMHzLrKMLbYTDgDBUizjC1/Hzrq6N3nN83GZQfLQW+ZlCC1p8zzN4/xsznxpm5d5za4UKHNeswF/0+h18ySKB5PNd5r56AJwQ4PkGOk2ryV1U40KuLrl9JZEbMXj7CKEd1OtzXGI2T/N1N1/xTDjcomsdyXHxwnGA0IXxPQvbWNku/yjP3WKnDpgdX9FiaoW5YazkohIjJpHIU8Hdivk+O19L1tekmkQjMSHSKjOp8BeMZjFL5ueFk5Q8DwfDDDcLpGF9zVJ8uUnhomdZrGRSAn82SP7BC4+ggjSN5hr68zMBdETYgFBm5OyOIUnYzGCOENJgj4pCN8MK2OijhzPD6FxNMcpQRPk4NaCLGsPIXRlHT8IuOkS8tsfiDAsGUx+XF2C/niV7OUv1mkdztLRQZY79YgDJQTFdqgPfIDQLGLMYRIn5s47yVsua6J9MlMrMK7jz7GeYlsuyhSgwEOIwQLANqdRZTDBaCmoabjgHI3dti6o65NY+zwFRPzuXBeTwxos5jNs4rvVHbUgdXwc3wQQocxzFCjQRnYUc/486uFJgTWjFyt7dZmHXkP9pk5eUB5h4Icbv8WkgMlKyPldqGG0usdLgOofbrRKea6Z5Mm1mYCrFXlSLiRYwRGiQ4go7rfi0Z0iCoDcFkQvEbdarf2kXhoQbJXIhfcJcmkHoAtvDBsDNizpLlJJ+8lLGtT5KQDG3ei08VqqtV5tdWjYF2B3My7xh6cgXLweIPC11R37p6MfAtM+LYk2E34maMP9KRk6RvsbAa4jL3kOM3DBJ0Q0o9nRVA/O8APKhlZN4dd1K6CDQhOefA23bHJji1wz0+gwcWmbabOKuDODvUDdUmDJrh00T9rS7yBI67aGIYU+T4NC08MS58X7LmVpQ6sNAJe3CT375ycQQUyLDMAhHPcZJzXXIuq9zaOFkVfs4QX6SBgAD1hNBtUxSsfy+KGCtEiJ/i+Z6VOLtZgdq3mkmZdPyFgGkGgH1kcdRJOiVqnwrSNgXnCRAZAiJeIuGwlTi9QVr6mtukevHsx9sEdcSLtHiTHB6RoA2hUH/vyeIYJSDmaRvm81bitE6sKkfClVi3Q9NFHpSQakgradPT6dC8FpDKxGlzpJ6RqEqkJm9qlme637vqfXD3VkBlXlCN11XhuGq0VUdaQKpQV4y0hFRNu7UyLa0gLXBOF9jT01ztyC67J9EsUxh7GWQaeJYmfwNe9Vne79rcBkxjDFIElpmjyWdsglMSYb+T4qoxCaALTKrCrXqD3XqDoi4wCaBzjGueM1qgoTo/03n2vuNXG1t1/hKBqnxWS/xOF7lvQ2t6xbaTtnOzAy1rRtQ7p5/4XrOLyevq3u+G3bCrZP8F9X0zAID2NJAAAAAASUVORK5CYII="
    )
    return _Resp(png, mimetype="image/png")


@word_constructor.get("/onlyoffice-plugin/index.html")
def oo_plugin_index():
    from flask import Response as _Resp
    html = """<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <link rel="stylesheet" href="styles.css">
  <!-- plugins.js sets up window.Asc.plugin — must load before script.js -->
  <script src="/onlyoffice/sdkjs-plugins/v1/plugins.js"></script>
</head>
<body>
  <div id="app">
    <div class="panel-header">
      <div class="search-wrap">
        <input id="search" type="text" placeholder="Поиск параметра…" autocomplete="off" spellcheck="false">
        <button id="btnRefresh" class="refresh-btn" title="Обновить">&#8635;</button>
      </div>
      <div class="mode-row">
        <label class="mode-label">
          <input type="radio" name="mode" value="placeholder" checked>
          <span>Метка <code>[…]</code></span>
        </label>
        <label class="mode-label">
          <input type="radio" name="mode" value="value">
          <span>Значение</span>
        </label>
      </div>
    </div>
    <div id="status" class="status" style="display:none"></div>
    <div id="list" class="param-list"></div>
  </div>
  <script src="script.js"></script>
</body>
</html>"""
    return _Resp(html, mimetype="text/html")


@word_constructor.get("/onlyoffice-plugin/script.js")
def oo_plugin_script():
    from flask import Response as _Resp
    js = r"""
/* KazUni — OnlyOffice sidebar plugin
   Rules:
   - window.Asc is injected by OnlyOffice AFTER the script loads.
     Nothing outside init() may touch window.Asc or the DOM.
   - session_id arrives via initData (config.json), with localStorage fallback.
   - Uses fetch + async/await throughout.                                     */

const API_BASE    = "/services/word-constructor";
const LS_SYNC_KEY = "oo_params_sync";

/* ── Module state (all set inside init) ── */
let sessionId  = "";
let paramsUrl  = "";
let allParams  = [];
let insertMode = "placeholder";
let debounceT  = null;

/* ── DOM refs (set inside init after DOMContentLoaded) ── */
let elList, elSearch, elStatus, elRefresh;

/* ═══════════════════════════════ API ═══════════════════════════════════ */

async function fetchParams() {
  if (!paramsUrl) {
    showStatus("Откройте шаблон через конструктор — сессия не определена.", "error");
    return;
  }
  showStatus("Загрузка…", "loading");
  try {
    const res = await fetch(paramsUrl);
    if (res.status === 404) {
      showStatus("Сессия не найдена или истекла.\nОткройте шаблон заново из 1С.", "error");
      return;
    }
    if (!res.ok) {
      showStatus(`Ошибка сервера (${res.status})`, "error");
      return;
    }
    const data = await res.json();
    allParams = data.params ?? [];
    if (allParams.length === 0) {
      showStatus("Параметры не добавлены.\nДобавьте их в левой панели конструктора.", "empty");
    } else {
      hideStatus();
      renderList();
    }
  } catch {
    showStatus("Нет связи с сервером", "error");
  }
}

/* ═══════════════════════════════ Render ════════════════════════════════ */

function renderList() {
  const q = elSearch.value.trim().toLowerCase();
  const items = q
    ? allParams.filter(p =>
        p.label.toLowerCase().includes(q) || p.key.toLowerCase().includes(q))
    : allParams;

  elList.innerHTML = "";

  if (items.length === 0) {
    elList.innerHTML = '<div class="empty-msg">Ничего не найдено</div>';
    return;
  }

  for (const p of items) {
    const el = document.createElement("div");
    el.className = "param-item";
    el.setAttribute("role", "button");
    el.setAttribute("tabindex", "0");
    el.innerHTML =
      `<div class="param-label">${esc(p.label)}</div>` +
      `<div class="param-key">[${esc(p.key)}]</div>` +
      (p.value ? `<div class="param-value">${esc(p.value)}</div>` : "");

    const doInsert = () => {
      const text = (insertMode === "value" && p.value) ? p.value : `[${p.key}]`;
      insertText(text, el);
    };
    el.addEventListener("click", doInsert);
    el.addEventListener("keydown", e => {
      if (e.key === "Enter" || e.key === " ") { e.preventDefault(); doInsert(); }
    });
    elList.appendChild(el);
  }
}

/* ═══════════════════════════════ Insert ════════════════════════════════ */

function insertText(text, el) {
  /* PasteText is the OnlyOffice plugin API — no clipboard involved */
  window.Asc.plugin.executeMethod("PasteText", [text], () => {});
  flash(el);
}

function flash(el) {
  el?.classList.add("inserted");
  setTimeout(() => el?.classList.remove("inserted"), 700);
}

/* ═══════════════════════════════ UI helpers ════════════════════════════ */

function showStatus(msg, type = "") {
  elStatus.className = `status ${type}`;
  elStatus.textContent = msg;
  elStatus.style.display = "block";
  elList.style.display   = "none";
}

function hideStatus() {
  elStatus.style.display = "none";
  elList.style.display   = "";
}

function esc(s) {
  return String(s ?? "")
    .replace(/&/g, "&amp;").replace(/</g, "&lt;")
    .replace(/>/g, "&gt;").replace(/"/g, "&quot;");
}

function debounce(fn, ms) {
  return (...args) => {
    clearTimeout(debounceT);
    debounceT = setTimeout(() => fn(...args), ms);
  };
}

/* ═══════════════════════════════ OnlyOffice lifecycle ══════════════════ */
/* window.Asc is NOT available when this script file loads —
   OnlyOffice injects it later and then calls window.Asc.plugin.init().
   We must not access window.Asc outside of the callbacks below.          */

window.Asc = window.Asc ?? {};
window.Asc.plugin = window.Asc.plugin ?? {};

window.Asc.plugin.init = function (data) {
  /* data = initData from config.json = the session_id */
  sessionId = (typeof data === "string" && data.trim()) ? data.trim() : "";
  if (!sessionId) {
    try { sessionId = localStorage.getItem("oo_session_id") ?? ""; } catch { /* ok */ }
  }
  paramsUrl = sessionId
    ? `${API_BASE}/api/template-builder/${sessionId}/params-list`
    : "";

  /* DOM is ready by the time init() is called */
  elList    = document.getElementById("list");
  elSearch  = document.getElementById("search");
  elStatus  = document.getElementById("status");
  elRefresh = document.getElementById("btnRefresh");

  elSearch.addEventListener("input", debounce(() => renderList(), 250));

  document.querySelectorAll('input[name="mode"]').forEach(r =>
    r.addEventListener("change", e => { insertMode = e.target.value; renderList(); })
  );

  elRefresh.addEventListener("click", () => fetchParams());

  /* Refresh when the template builder sidebar adds/removes params */
  window.addEventListener("storage", e => {
    if (e.key === LS_SYNC_KEY) fetchParams();
  });

  fetchParams();
};

window.Asc.plugin.button = () => {};
"""
    return _Resp(js, mimetype="application/javascript")


@word_constructor.get("/onlyoffice-plugin/styles.css")
def oo_plugin_styles():
    from flask import Response as _Resp
    css = """
* { box-sizing: border-box; margin: 0; padding: 0; }

body {
  font-family: "Segoe UI", Tahoma, Arial, sans-serif;
  font-size: 13px;
  color: #1a2e42;
  background: #f4f7fa;
  height: 100vh;
  overflow: hidden;
  display: flex;
  flex-direction: column;
}

#app {
  display: flex;
  flex-direction: column;
  height: 100%;
  overflow: hidden;
}

/* ── Header ── */
.panel-header {
  background: #fff;
  border-bottom: 1px solid #d6dde5;
  padding: 10px 10px 8px;
  flex-shrink: 0;
}


/* ── Search ── */
.search-wrap {
  display: flex;
  gap: 4px;
  margin-bottom: 8px;
}

#search {
  flex: 1;
  padding: 6px 8px;
  border: 1px solid #d6dde5;
  border-radius: 6px;
  font-size: 12px;
  background: #f8fafc;
  outline: none;
  color: #1a2e42;
}
#search:focus { border-color: #1f4b7a; background: #fff; }

.refresh-btn {
  padding: 0 8px;
  border: 1px solid #d6dde5;
  border-radius: 6px;
  background: #f8fafc;
  color: #5c6b7a;
  font-size: 15px;
  cursor: pointer;
  line-height: 1;
}
.refresh-btn:hover { background: #e8f0f8; color: #1f4b7a; }

/* ── Mode toggle ── */
.mode-row {
  display: flex;
  gap: 12px;
}

.mode-label {
  display: flex;
  align-items: center;
  gap: 4px;
  cursor: pointer;
  font-size: 11px;
  color: #5c6b7a;
  user-select: none;
}
.mode-label input { cursor: pointer; }
.mode-label code {
  font-size: 10px;
  background: #eef2f6;
  padding: 1px 4px;
  border-radius: 3px;
}

/* ── Status messages ── */
.status {
  padding: 16px 12px;
  text-align: center;
  font-size: 12px;
  color: #5c6b7a;
}
.status.error   { color: #b42318; }
.status.loading { color: #1f4b7a; }
.status.empty   { color: #8a9bac; }

/* ── Param list ── */
.param-list {
  flex: 1;
  overflow-y: auto;
  padding: 6px 8px 8px;
}

.param-list::-webkit-scrollbar       { width: 5px; }
.param-list::-webkit-scrollbar-track { background: transparent; }
.param-list::-webkit-scrollbar-thumb { background: #c8d4de; border-radius: 3px; }

.param-item {
  padding: 8px 10px;
  border-radius: 8px;
  border: 1px solid #e4eaf0;
  background: #fff;
  margin-bottom: 5px;
  cursor: pointer;
  transition: background 100ms, border-color 100ms, transform 80ms;
  outline: none;
}
.param-item:hover  { background: #e8f0f8; border-color: #b3cce8; }
.param-item:focus  { border-color: #1f4b7a; box-shadow: 0 0 0 2px rgba(31,75,122,.15); }
.param-item.inserted {
  background: #d4edda;
  border-color: #28a745;
  transform: scale(.98);
}

.param-label {
  font-weight: 600;
  font-size: 12px;
  color: #1a2e42;
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}

.param-key {
  font-family: "Consolas", "Courier New", monospace;
  font-size: 10px;
  color: #1f4b7a;
  margin-top: 2px;
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}

.param-value {
  font-size: 11px;
  color: #5c6b7a;
  margin-top: 2px;
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}

.empty-msg {
  text-align: center;
  color: #8a9bac;
  font-size: 12px;
  padding: 20px 0;
}
"""
    return _Resp(css, mimetype="text/css")


# ── Params-list endpoint: used by the OO plugin to fetch session params ──
@word_constructor.get("/api/template-builder/<session_id>/params-list")
def api_template_builder_params_list(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        return jsonify({"error": "not found"}), 404
    raw = meta.get("params", [])
    return jsonify({
        "params": [{"key": p, "label": p, "value": ""} for p in raw if p]
    })


# WebSocket endpoint is registered in the main app.py via flask-sock
# (Werkzeug 3.x requires flask-sock for WebSocket routes; Blueprint GET routes
#  return 400 for WebSocket upgrade requests).


class _HtmlToDocxBuilder(HTMLParser):
    """Convert a simple editable HTML fragment to a python-docx Document."""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.para = None
        self._bold = 0
        self._italic = 0
        self._under = 0
        self._list_type = []  # type: list[str]  stack: "ol" | "ul"
        self._heading = None  # type: Optional[str]

    # ------------------------------------------------------------------
    def _new_para(self, style=None, align=None):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        self.para = self.doc.add_paragraph(style=style)
        if align:
            _map = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if align in _map:
                self.para.alignment = _map[align]

    def _attrs_dict(self, attrs):
        return dict(attrs)

    def _align_from_attrs(self, attrs):
        d = self._attrs_dict(attrs)
        cls = d.get("class", "")
        if "ql-align-center" in cls:
            return "center"
        if "ql-align-right" in cls:
            return "right"
        if "ql-align-justify" in cls:
            return "justify"
        return None

    def _image_width_from_attrs(self, attrs):
        from docx.shared import Inches

        d = self._attrs_dict(attrs)
        width_attr = str(d.get("width", "")).strip()
        style_attr = str(d.get("style", "")).strip()
        px_match = None

        if width_attr.isdigit():
            px_match = int(width_attr)
        else:
            m = re.search(r"width\s*:\s*([0-9]+)px", style_attr, re.I)
            if m:
                px_match = int(m.group(1))

        if not px_match:
            return None
        return Inches(px_match / 96.0)

    def _insert_data_image(self, attrs):
        d = self._attrs_dict(attrs)
        src = str(d.get("src", "")).strip()
        if not src.startswith("data:image/") or ";base64," not in src:
            return

        try:
            _, encoded = src.split(";base64,", 1)
            image_bytes = base64.b64decode(encoded)
        except Exception:
            return

        if self.para is None:
            self._new_para()

        run = self.para.add_run()
        width = self._image_width_from_attrs(attrs)
        try:
            if width is not None:
                run.add_picture(BytesIO(image_bytes), width=width)
            else:
                run.add_picture(BytesIO(image_bytes))
        except Exception:
            return

    # ------------------------------------------------------------------
    def handle_starttag(self, tag, attrs):
        align = self._align_from_attrs(attrs)
        if tag == "p":
            self._new_para(align=align)
        elif tag in ("h1", "h2", "h3"):
            level = tag[1]
            self._heading = tag
            try:
                self._new_para(style=f"Heading {level}", align=align)
            except Exception:
                self._new_para(align=align)
        elif tag in ("ol", "ul"):
            self._list_type.append(tag)
        elif tag == "li":
            style = "List Number" if (self._list_type and self._list_type[-1] == "ol") else "List Bullet"
            try:
                self._new_para(style=style, align=align)
            except Exception:
                self._new_para(align=align)
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag == "u":
            self._under += 1
        elif tag == "br":
            if self.para is not None:
                self.para.add_run("\n")
            else:
                self._new_para()
                self.para.add_run("\n")
        elif tag == "img":
            self._insert_data_image(attrs)

    def handle_endtag(self, tag):
        if tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag == "u":
            self._under = max(0, self._under - 1)
        elif tag in ("ol", "ul"):
            if self._list_type:
                self._list_type.pop()
        elif tag in ("h1", "h2", "h3"):
            self._heading = None

    def handle_data(self, data):
        if not data:
            return
        if self.para is None:
            self._new_para()
        run = self.para.add_run(data)
        run.bold = self._bold > 0
        run.italic = self._italic > 0
        run.underline = self._under > 0


def _html_to_docx(html: str) -> Document:
    """Convert Quill HTML to a python-docx Document."""
    doc = Document()
    # Remove the default empty paragraph python-docx adds
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    builder = _HtmlToDocxBuilder(doc)
    builder.feed(html)
    # Ensure at least one paragraph
    if not doc.paragraphs:
        doc.add_paragraph("")
    return doc


@word_constructor.post("/api/template/build")
def api_template_build():
    """
    Convert editor content into a .docx template file.
    Body (preferred): {"html": "<p>...</p>", "filename": "template.docx"}
    Body (legacy):    {"lines": ["paragraph text", ...], "filename": "template.docx"}
    [ParamName] markers in text are preserved as word-constructor placeholders.
    """
    data = request.get_json(silent=True) or {}
    download_name = str(data.get("filename", "template.docx")).strip() or "template.docx"
    if not download_name.lower().endswith(".docx"):
        download_name += ".docx"

    html = data.get("html", "")
    if html and str(html).strip():
        doc = _html_to_docx(str(html))
    else:
        lines = data.get("lines", [])
        if not isinstance(lines, list):
            return jsonify({"error": "Expected 'lines' array or 'html' string"}), 400
        if not any(str(ln).strip() for ln in lines):
            return jsonify({"error": "Document is empty"}), 400

        doc = Document()
        if doc.paragraphs:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        for line in lines:
            doc.add_paragraph(str(line))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=download_name,
        mimetype=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        etag=False,
        conditional=False,
    )


_JSON_REPLACE_RESERVED_KEYS = {
    "filename",
    "name",
    "content_base64",
    "base64",
    "content",
    "params",
    "placeholders",
    "values",
    "UseAI",
    "use_ai",
    "ИспользоватьAI",
    "PromtAI",
    "PromptAI",
    "prompt_ai",
}


def _replace_values_from_json_payload(payload: dict[str, Any]) -> dict[str, Any]:
    parsed = payload.get("params", payload.get("placeholders", payload.get("values")))
    if parsed is None:
        parsed = {}
    if not isinstance(parsed, dict):
        raise ValueError("'params'/'placeholders' must be a JSON object")

    values = dict(parsed)
    for key, value in payload.items():
        if str(key) in _JSON_REPLACE_RESERVED_KEYS:
            continue
        values.setdefault(str(key), value)
    return values


def _parse_replace_payload() -> tuple[str, bytes, dict[str, str], dict[str, list], dict[str, list[dict[str, str]]]]:
    if request.is_json:
        payload = request.get_json(silent=True) or {}
        filename = str(payload.get("filename") or payload.get("name") or "document.docx")
        content_base64 = str(
            payload.get("content_base64")
            or payload.get("base64")
            or payload.get("content")
            or ""
        ).strip()
        if not content_base64:
            raise ValueError("Missing 'content_base64'")
        content = _safe_b64decode(content_base64)
        parsed = _replace_values_from_json_payload(payload)
    else:
        upload = (
            request.files.get("template")
            or request.files.get("document")
            or request.files.get("file")
        )
        if not upload:
            raise ValueError("Missing .docx file field: use 'template', 'document', or 'file'")
        filename = upload.filename or "document.docx"
        content = upload.read()
        raw = request.form.get("params") or request.form.get("placeholders") or "{}"
        parsed = json.loads(raw) if raw.strip() else {}

    if not content:
        raise ValueError("Uploaded document is empty")
    if not isinstance(parsed, dict):
        raise ValueError("'params'/'placeholders' must be a JSON object")

    slot_values: dict[str, str] = {}
    table_params: dict[str, list] = {}
    table_object_params: dict[str, list[dict[str, str]]] = {}
    for key, value in parsed.items():
        if isinstance(value, list):
            object_rows = _normalize_table_object_rows(value)
            if object_rows:
                table_object_params[str(key)] = object_rows
            else:
                table_params[str(key)] = [
                    [str(cell) for cell in row]
                    for row in value
                    if isinstance(row, list)
                ]
        else:
            slot_values[str(key)] = str(value)

    if not filename.lower().endswith(".docx"):
        filename = f"{Path(filename).stem or 'document'}.docx"

    return filename, content, slot_values, table_params, table_object_params


@word_constructor.post("/api/replace")
@word_constructor.post("/api/1c/replace")
def api_replace_docx_placeholders():
    """
    Stateless 1C endpoint: replace placeholders in a Word .docx and return the
    resulting .docx immediately.

    Multipart:
      - template/document/file: .docx
      - params/placeholders: JSON object {"Key": "Value"}

    JSON:
      - filename: "document.docx"
      - content_base64: "<base64 docx>"
      - params/placeholders: {"Key": "Value"}
    """
    try:
        filename, template_bytes, slot_values, table_params, table_object_params = _parse_replace_payload()
        doc = Document(BytesIO(template_bytes))
        log_key = f"replace-{uuid.uuid4().hex}"
        slot_values, slot_occurrence_values, _ = _ai_correct_slot_values(
            doc, slot_values, "",
            log_key=log_key,
            call_log={"document_name": filename},
        )
        result_bytes = fill_docx(
            template_bytes, slot_values, table_params,
            table_object_params=table_object_params,
            slot_occurrence_values=slot_occurrence_values,
        )
    except json.JSONDecodeError as exc:
        return jsonify({"error": f"Invalid JSON params: {exc}"}), 400
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"Cannot replace placeholders: {exc}"}), 400

    return send_file(
        BytesIO(result_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        etag=False,
        conditional=False,
    )


def _create_onlyoffice_edit_session(filename: str, document_bytes: bytes, params: list[str] | None = None) -> dict[str, Any]:
    session_id = str(uuid.uuid4())
    sdir = _session_dir(session_id)
    sdir.mkdir(parents=True, exist_ok=True)

    expires_at = time.time() + SESSION_TB_TTL_SECONDS
    meta = {
        "id": session_id,
        "type": "template_builder",
        "workflow": "manual_edit",
        "params": params or [],
        "filename": filename,
        "status": "editing",
        "last_saved_at": time.time(),
        "expires_at": expires_at,
        "expires_at_iso": datetime.fromtimestamp(expires_at, timezone.utc).isoformat(),
    }
    if getattr(g, "api_client_id", None):
        meta["client_id"] = g.api_client_id
    _write_meta(session_id, meta)
    _session_template_path(session_id).write_bytes(document_bytes)
    return meta


@word_constructor.post("/api/1c/replace-edit")
@word_constructor.post("/api/replace-edit")
def api_replace_and_open_edit_session():
    """
    Replace placeholders, store the resulting Word document in an OnlyOffice edit
    session, and return URLs for user editing plus 1C refresh/download.
    """
    try:
        filename, template_bytes, slot_values, table_params, table_object_params = _parse_replace_payload()
        use_ai, prompt_ai = _parse_ai_replace_options()
        doc = Document(BytesIO(template_bytes))
        slot_occurrence_values: dict[tuple[str, int], str] = {}
        use_ai_log_key = f"useai-{uuid.uuid4().hex}" if use_ai else None
        use_ai_log: dict[str, Any] | None = {"key": use_ai_log_key, "document_name": filename} if use_ai_log_key else None
        review_summary = ""
        if use_ai:
            slot_values, slot_occurrence_values, review_summary = _ai_correct_slot_values(
                doc,
                slot_values,
                prompt_ai,
                use_ai_log_key,
                use_ai_log,
            )
        result_bytes = fill_docx(
            template_bytes,
            slot_values,
            table_params,
            table_object_params=table_object_params,
            slot_occurrence_values=slot_occurrence_values,
        )
        meta = _create_onlyoffice_edit_session(filename, result_bytes, sorted(slot_values.keys()))
    except json.JSONDecodeError as exc:
        return jsonify({"error": f"Invalid JSON params: {exc}"}), 400
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"Cannot create edit session: {exc}"}), 400

    session_id = meta["id"]
    base = "/services/word-constructor"
    response_payload = {
        "id": session_id,
        "editor_url": f"{base}/template-builder/{session_id}?source=from1c",
        "status_url": f"{base}/api/1c/edit-sessions/{session_id}/status",
        "update_url": f"{base}/api/1c/edit-sessions/{session_id}/document",
        "download_url": f"{base}/api/1c/edit-sessions/{session_id}/document?forcesave=0",
        "expires_at": meta["expires_at_iso"],
    }
    if use_ai_log_key:
        response_payload["use_ai_log_key"] = use_ai_log_key
    if review_summary:
        response_payload["review_summary"] = review_summary
    return jsonify(response_payload), 201


@word_constructor.get("/api/1c/edit-sessions/<session_id>/status")
def api_1c_edit_session_status(session_id: str):
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        return jsonify({"status": "not_found", "id": session_id}), 404
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"status": "expired", "id": session_id}), 410
    status = meta.get("status", "editing")
    response = {
        "id": session_id,
        "status": status,
        "last_saved_at": meta.get("last_saved_at"),
        "expires_at": meta.get("expires_at_iso"),
    }
    if status == "ready":
        response["download_url"] = f"/services/word-constructor/api/1c/edit-sessions/{session_id}/document?forcesave=0"
        response["update_url"] = response["download_url"]
    return jsonify(response)


def _force_save_session_document(session_id: str, meta: dict) -> tuple[bool, dict[str, Any] | None, str | None]:
    path = _session_template_path(session_id)
    if not path.exists():
        return False, None, "Document file not found"

    editor_key = meta.get("editor_key")
    if not editor_key:
        return True, {"skipped": True, "reason": "OnlyOffice editor was not opened"}, None

    key = str(editor_key)
    previous_saved_at = meta.get("last_saved_at")
    try:
        result = _builder_forcesave(session_id, key)
    except Exception as exc:
        return False, None, f"Force save failed: {exc}"

    error_code = int(result.get("error", 0) or 0)
    if error_code == 0:
        if not _wait_for_builder_save(session_id, previous_saved_at):
            return False, result, "Force save timed out waiting for callback"
    elif error_code != 4:
        return False, result, "ONLYOFFICE rejected force save"

    return True, result, None


@word_constructor.get("/api/1c/edit-sessions/<session_id>/document")
@word_constructor.post("/api/1c/edit-sessions/<session_id>/document")
def api_1c_edit_session_document(session_id: str):
    """
    1C "Обновить" endpoint: force-save the open OnlyOffice editor and return the
    latest saved .docx without deleting the session.
    """
    meta = _read_meta(session_id)
    if meta is None or meta.get("type") != "template_builder":
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    force = request.args.get("forcesave", "1") != "0"
    if force:
        ok, result, error = _force_save_session_document(session_id, meta)
        if not ok:
            return jsonify({"error": error, "details": result}), 502

    path = _session_template_path(session_id)
    if not path.exists():
        abort(404)

    filename = meta.get("filename", "document.docx")
    should_close = meta.get("workflow") == "manual_edit" and meta.get("status") == "ready"
    if should_close:
        file_bytes = path.read_bytes()
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return _send_file_compat(
            BytesIO(file_bytes),
            as_attachment=True,
            download_name=filename,
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
        )

    return _send_file_compat(
        path,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    )


@word_constructor.post("/api/create")
def api_create():
    """
    1C endpoint: POST multipart/form-data
      - template: .docx file  (field name must be 'template')
      - params:   JSON string {"key": "value", ...}
                  OR omit 'params' and send each key as its own form field

    Returns: {"id": "...", "editor_url": "...", "expires_at": "..."}
    """
    # ── locate the template file ─────────────────────────────────────────────
    template_file = request.files.get("template")
    if not template_file:
        received_fields = list(request.files.keys()) or ["(none)"]
        received_form  = list(request.form.keys()) or ["(none)"]
        return jsonify({
            "error": "Missing 'template' file field",
            "hint": (
                "Send as multipart/form-data with the .docx file "
                "in a field named 'template'."
            ),
            "received_files":  received_fields,
            "received_form":   received_form,
        }), 400

    # ── parse params ─────────────────────────────────────────────────────────
    # Strategy 1: 'params' field contains a JSON string  {"key": "value"}
    # Strategy 2: 'params' field absent — use all non-reserved form fields
    RESERVED = {"template", "params"}
    params: dict[str, str] = {}

    if "params" in request.form:
        params_raw = request.form["params"].strip()
        if not params_raw or params_raw == "null":
            parsed = {}
        else:
            try:
                parsed = json.loads(params_raw)
                if not isinstance(parsed, dict):
                    return jsonify({
                        "error": "Field 'params' must be a JSON object, not an array or scalar",
                        "received": params_raw[:200],
                        "example": '{"ФИО": "Иванов А.А.", "Таблица": [["Кол1","Кол2"],["A","B"]]}',
                    }), 400
            except json.JSONDecodeError as exc:
                return jsonify({
                    "error": f"Field 'params' is not valid JSON: {exc}",
                    "received": params_raw[:200],
                    "example": '{"ФИО": "Иванов А.А.", "Должность": "Директор"}',
                }), 400
    else:
        # Fallback: treat every other form field as a param key (string params only)
        parsed = {
            k: request.form[k]
            for k in request.form
            if k not in RESERVED
        }

    # Split string params vs table params (list values)
    params: dict[str, str] = {}
    table_params: dict[str, list] = {}
    for k, v in parsed.items():
        if isinstance(v, list):
            table_params[str(k)] = [[str(c) for c in row] for row in v if isinstance(row, list)]
        else:
            params[str(k)] = str(v)

    # ── validate the docx ────────────────────────────────────────────────────
    template_bytes = template_file.read()
    if not template_bytes:
        return jsonify({"error": "Uploaded 'template' file is empty"}), 400
    try:
        doc = Document(BytesIO(template_bytes))
    except Exception:
        return jsonify({
            "error": "Cannot read 'template' as a .docx file",
            "hint": (
                "Only Word Open XML (.docx) files are accepted. "
                ".doc, .odt, .pdf and other formats are not supported."
            ),
            "filename": template_file.filename or "(unknown)",
        }), 400

    template_keys = _extract_placeholder_keys(doc)

    session_id = str(uuid.uuid4())
    sdir = _session_dir(session_id)
    sdir.mkdir(parents=True, exist_ok=True)
    (sdir / "template.docx").write_bytes(template_bytes)

    expires_at = time.time() + SESSION_TTL_SECONDS
    meta = {
        "id": session_id,
        "params": params,
        "table_params": table_params,
        "slot_values": dict(params),   # editable copy; updated as user types
        "template_keys": template_keys,
        "filename": template_file.filename or "document.docx",
        "expires_at": expires_at,
        "expires_at_iso": datetime.fromtimestamp(expires_at, timezone.utc).isoformat(),
        "params_version": 1,
    }
    _write_meta(session_id, meta)

    return jsonify({
        "id": session_id,
        "editor_url": f"/services/word-constructor/editor/{session_id}",
        "expires_at": meta["expires_at_iso"],
    })


@word_constructor.post("/api/sessions/<session_id>/params")
def api_update_params(session_id: str):
    """
    1C endpoint: update / add params to an existing session.
    Only the params list is updated — the open editor document is NOT re-rendered.

    Body: JSON {"key": "value", ...}   OR  form field params=<JSON>
    Returns: {"params_version": N, "params": {...}}
    """
    meta = _read_meta(session_id)
    if meta is None:
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    # Accept JSON body or form field
    if request.is_json:
        new_params = request.get_json(silent=True) or {}
    else:
        raw = request.form.get("params", "{}")
        try:
            new_params = json.loads(raw)
        except json.JSONDecodeError:
            return jsonify({"error": "Invalid JSON"}), 400

    if not isinstance(new_params, dict):
        return jsonify({"error": "Expected JSON object"}), 400

    if "table_params" not in meta:
        meta["table_params"] = {}

    for k, v in new_params.items():
        if isinstance(v, list):
            meta["table_params"][str(k)] = [[str(c) for c in row] for row in v if isinstance(row, list)]
        else:
            meta["params"][str(k)] = str(v)

    meta["params_version"] = meta.get("params_version", 1) + 1
    _write_meta(session_id, meta)

    return jsonify({
        "params_version": meta["params_version"],
        "params": meta["params"],
        "table_params": {k: len(v) for k, v in meta["table_params"].items()},
    })


@word_constructor.get("/api/sessions/<session_id>/params")
def api_get_params(session_id: str):
    """
    Polling endpoint for the UI: returns current params + version.
    The UI polls this to pick up param updates pushed by 1C.
    """
    meta = _read_meta(session_id)
    if meta is None:
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    return jsonify({
        "params_version": meta.get("params_version", 1),
        "params": meta["params"],
    })


@word_constructor.get("/editor/<session_id>")
def editor(session_id: str):
    meta = _read_meta(session_id)
    if meta is None:
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        abort(410)

    template_bytes = (_session_dir(session_id) / "template.docx").read_bytes()
    doc = Document(BytesIO(template_bytes))
    table_params = meta.get("table_params", {})
    # Use slot_values (user edits) so refreshing the editor preserves changes;
    # fall back to original params if never edited.
    render_params = meta.get("slot_values") or meta["params"]
    doc_html = docx_to_html(doc, render_params, table_params)

    return render_template(
        "word_constructor/editor.html",
        session_id=session_id,
        params=meta["params"],
        table_params=table_params,
        params_version=meta.get("params_version", 1),
        doc_html=doc_html,
        filename=meta.get("filename", "document.docx"),
        expires_at=meta["expires_at_iso"],
    )


@word_constructor.post("/api/sessions/<session_id>/inject")
def api_inject_table(session_id: str):
    """
    Browser endpoint: save the list of tables dragged into the document.
    Body: {"injected_tables": [{"body_idx": N, "key": "TableName"}, ...]}
    """
    meta = _read_meta(session_id)
    if meta is None:
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    data = request.get_json(silent=True) or {}
    injected = data.get("injected_tables")
    if isinstance(injected, list):
        meta["injected_tables"] = [
            {"body_idx": int(i.get("body_idx", 0)), "key": str(i.get("key", ""))}
            for i in injected
            if isinstance(i, dict)
        ]
        _write_meta(session_id, meta)
    return jsonify({"ok": True, "count": len(meta.get("injected_tables", []))})


@word_constructor.post("/api/sessions/<session_id>/slots")
def api_save_slots(session_id: str):
    """
    Browser auto-save: persist current slot values edited by the user.
    Body: JSON {"key": "edited_value", ...}
    """
    meta = _read_meta(session_id)
    if meta is None:
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    updates = request.get_json(silent=True) or {}
    if not isinstance(updates, dict):
        return jsonify({"error": "Expected JSON object"}), 400

    if "slot_values" not in meta:
        meta["slot_values"] = _meta_string_params(meta)
    meta["slot_values"].update({str(k): str(v) for k, v in updates.items()})
    _write_meta(session_id, meta)
    return jsonify({"saved": len(updates)})


@word_constructor.get("/api/download/<session_id>")
def api_download(session_id: str):
    """
    GET — generate and return the filled .docx using saved slot values, then delete the session.
    No request body needed; values were auto-saved by the editor as the user typed.
    """
    meta = _read_meta(session_id)
    if meta is None:
        abort(404)
    if _is_expired(meta):
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return jsonify({"error": "Session expired"}), 410

    if meta.get("type") == "template_builder":
        builder_html = str(meta.get("builder_html", "") or "").strip()
        if not builder_html:
            return jsonify({"error": "No saved builder content"}), 409
        try:
            doc = _html_to_docx(builder_html)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = meta.get("filename", "template.docx")
        shutil.rmtree(_session_dir(session_id), ignore_errors=True)
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            etag=False,
            conditional=False,
        )

    # Use saved edits; fall back to original params if slot_values not yet stored
    slot_values = meta.get("slot_values") or _meta_string_params(meta)

    template_bytes = (_session_dir(session_id) / "template.docx").read_bytes()
    try:
        result_bytes = fill_docx(
            template_bytes,
            slot_values,
            meta.get("table_params", {}),
            meta.get("injected_tables", []),
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    shutil.rmtree(_session_dir(session_id), ignore_errors=True)

    filename = meta.get("filename", "document.docx")
    return send_file(
        BytesIO(result_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        etag=False,
        conditional=False,
    )


# ---------------------------------------------------------------------------
# Transform endpoints
# ---------------------------------------------------------------------------

@word_constructor.get("/api/transforms")
def api_transforms():
    """
    GET /api/transforms?value=<text>
    Returns all applicable transforms for a value with pre-computed results.
    Used by the editor UI to populate the fx panel.
    """
    value = request.args.get("value", "")
    return jsonify(get_transforms(value))


@word_constructor.post("/api/transform")
def api_transform():
    """
    POST /api/transform  {value, function}
    Apply a single named transform to a value.
    """
    data = request.get_json(silent=True) or {}
    value = str(data.get("value", ""))
    fn    = str(data.get("function", ""))
    if not fn:
        return jsonify({"error": "Missing 'function'"}), 400
    return jsonify({"result": apply_transform(fn, value)})
