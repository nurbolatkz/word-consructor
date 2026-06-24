"""Extraction — walk a .docx and find placeholder occurrences. No preprocessing."""
from __future__ import annotations

import re
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .types import Occurrence, SanityCheck, TextUnit

PLACEHOLDER_RE = re.compile(r"\{\{([^{}\n\r]{1,120})\}\}|\[([^\[\]\n\r]{1,120})\]")
_CONTEXT_CHARS_DEFAULT = 240


def match_key(match: re.Match) -> str:
    return (match.group(1) or match.group(2)).strip()


def para_full_text(para) -> str:
    return "".join(run.text for run in para.runs)


def cell_text(cell) -> str:
    parts: list[str] = []
    for para in cell.paragraphs:
        text = para_full_text(para)
        if text:
            parts.append(text)
    return "\n".join(parts)


def iter_text_units(doc: Document, include_headers_footers: bool = False) -> list[TextUnit]:
    units: list[TextUnit] = []
    paragraph_idx = 0
    table_idx = 0
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para = Paragraph(child, doc)
            units.append(TextUnit("body_paragraph", f"paragraph[{paragraph_idx}]", para_full_text(para)))
            paragraph_idx += 1
        elif tag == "tbl":
            table = Table(child, doc)
            for row_idx, row in enumerate(table.rows):
                row_cell_texts = tuple(cell_text(row_cell) for row_cell in row.cells)
                for cell_idx, _cell in enumerate(row.cells):
                    units.append(TextUnit(
                        "table_cell",
                        f"table[{table_idx}].row[{row_idx}].cell[{cell_idx}]",
                        row_cell_texts[cell_idx],
                        table_index=table_idx,
                        row_index=row_idx,
                        cell_index=cell_idx,
                        row_cell_texts=row_cell_texts,
                    ))
            table_idx += 1

    if not include_headers_footers:
        return units

    for section_idx, section in enumerate(doc.sections):
        for source_type, part in (("header", section.header), ("footer", section.footer)):
            paragraphs = [para_full_text(para) for para in part.paragraphs if para_full_text(para).strip()]
            for table in part.tables:
                for row in table.rows:
                    for c in row.cells:
                        text = cell_text(c)
                        if text.strip():
                            paragraphs.append(text)
            text = "\n".join(paragraphs)
            if text.strip():
                units.append(TextUnit(source_type, f"{source_type}[section={section_idx}]", text))
    return units


def document_plain_text(doc: Document) -> str:
    blocks: list[str] = []
    for unit in iter_text_units(doc):
        if unit.text.strip():
            blocks.append(f"[{unit.source_type} {unit.source_path}]\n{unit.text}")
    return "\n".join(blocks)


def document_full_text(doc: Document) -> str:
    """Full document text with structural labels — sent to AI as context."""
    blocks: list[str] = []
    for section_idx, section in enumerate(doc.sections):
        for part_name, part in (("header", section.header), ("footer", section.footer)):
            lines: list[str] = []
            for para in part.paragraphs:
                t = para_full_text(para).strip()
                if t:
                    lines.append(t)
            for table in part.tables:
                for row in table.rows:
                    for c in row.cells:
                        t = cell_text(c).strip()
                        if t:
                            lines.append(t)
            if lines:
                blocks.append(f"[{part_name} section={section_idx}]\n" + "\n".join(lines))
    for unit in iter_text_units(doc, include_headers_footers=False):
        if unit.text.strip():
            blocks.append(f"[{unit.source_type} {unit.source_path}]\n{unit.text}")
    return "\n\n".join(blocks)


def document_placeholder_scan_text(doc: Document) -> str:
    return "\n".join(unit.text for unit in iter_text_units(doc) if unit.text.strip())


def context_snippet(text: str, match: re.Match, window: int = _CONTEXT_CHARS_DEFAULT) -> str:
    start = max(0, match.start() - window)
    end = min(len(text), match.end() + window)
    snippet = re.sub(r"\s+", " ", text[start:end]).strip()
    if start > 0:
        snippet = "..." + snippet
    if end < len(text):
        snippet += "..."
    return snippet


def extract_placeholder_occurrences(doc: Document, slot_values: dict[str, str]) -> list[dict[str, Any]]:
    """Walk the document and return one dict per placeholder occurrence."""
    if not slot_values:
        return []
    wanted = set(slot_values)
    counts: dict[str, int] = {}
    result: list[dict[str, Any]] = []
    for unit in iter_text_units(doc):
        if not unit.text:
            continue
        for match in PLACEHOLDER_RE.finditer(unit.text):
            key = match_key(match)
            if key not in wanted:
                continue
            occ_idx = counts.get(key, 0)
            counts[key] = occ_idx + 1
            result.append({
                "id": f"{key}#{occ_idx}",
                "key": key,
                "placeholder": key,
                "occurrence": occ_idx + 1,
                "occurrence_index": occ_idx,
                "value": str(slot_values[key]),
                "original_value": str(slot_values[key]),
                "context": context_snippet(unit.text, match),
                "context_text": context_snippet(unit.text, match),
                "source_type": unit.source_type,
                "source_path": unit.source_path,
            })
    return result


def raw_placeholder_matches_from_doc(doc: Document, slot_values: dict[str, Any]) -> list[dict[str, Any]]:
    matches: list[dict[str, Any]] = []
    wanted = set(slot_values)
    for unit in iter_text_units(doc):
        if not unit.text:
            continue
        for match in PLACEHOLDER_RE.finditer(unit.text):
            key = match_key(match)
            if key not in wanted:
                continue
            matches.append({
                "placeholder": key,
                "source_type": unit.source_type,
                "source_path": unit.source_path,
                "context_text": context_snippet(unit.text, match),
            })
    return matches


def sanity_check_occurrence_counts(doc: Document, slot_values: dict[str, Any], occurrences: list[dict[str, Any]]) -> SanityCheck:
    wanted = set(slot_values)
    scan_text = document_placeholder_scan_text(doc)
    full_text_raw_count = sum(1 for match in PLACEHOLDER_RE.finditer(scan_text) if match_key(match) in wanted)
    raw_matches = raw_placeholder_matches_from_doc(doc, slot_values)
    raw_count = len(raw_matches)
    occurrence_count = len(occurrences)
    return SanityCheck(full_text_raw_count == occurrence_count and raw_count == occurrence_count, full_text_raw_count, raw_count, occurrence_count, raw_matches)


def extract_placeholder_contexts(doc: Document, slot_values: dict[str, str], max_snippets: int = 5) -> dict[str, list[str]]:
    contexts: dict[str, list[str]] = {key: [] for key in slot_values}
    wanted = set(slot_values)
    for unit in iter_text_units(doc):
        for match in PLACEHOLDER_RE.finditer(unit.text or ""):
            key = match_key(match)
            if key not in wanted:
                continue
            snippets = contexts.setdefault(key, [])
            if len(snippets) >= max_snippets:
                continue
            snippet = context_snippet(unit.text, match)
            if snippet and snippet not in snippets:
                snippets.append(snippet)
    return contexts


# Public notebook API -------------------------------------------------------

def walk_document(doc) -> list[TextUnit]:
    """Walk body paragraphs, every table cell's paragraphs, headers, footers."""
    return iter_text_units(doc, include_headers_footers=True)


def find_occurrences(text_units: list[TextUnit], placeholders: dict[str, str]) -> list[Occurrence]:
    """Regex-scan TextUnits for [PlaceholderName] matches present in placeholders."""
    occurrences: list[Occurrence] = []
    counts: dict[str, int] = {}
    wanted = set(placeholders)
    for unit in text_units:
        for match in PLACEHOLDER_RE.finditer(unit.text or ""):
            key = match_key(match)
            if key not in wanted:
                continue
            index = counts.get(key, 0)
            counts[key] = index + 1
            context = context_snippet(unit.text, match)
            occurrences.append(Occurrence(
                placeholder=key,
                occurrence_index=index,
                original_value=str(placeholders[key]),
                source_type=unit.source_type,
                source_path=unit.source_path,
                context_text=context,
                context=context,
                context_with_value=context.replace(match.group(0), str(placeholders[key]), 1),
                literal_placeholder=match.group(0),
            ))
    return occurrences


def sanity_check(text_units: list[TextUnit], occurrences: list[Occurrence]) -> tuple[bool, list[str]]:
    """Compare raw placeholder regex matches across all text_units against len(occurrences)."""
    raw_count = sum(1 for unit in text_units for _ in PLACEHOLDER_RE.finditer(unit.text or ""))
    occurrence_count = len(occurrences)
    if raw_count == occurrence_count:
        return True, []
    return False, [f"raw placeholder regex matches={raw_count}, occurrences={occurrence_count}"]
