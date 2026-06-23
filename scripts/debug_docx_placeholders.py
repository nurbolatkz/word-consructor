from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from word_constructor.app import (  # noqa: E402
    _PLACEHOLDER_RE,
    _document_placeholder_scan_text,
    _extract_placeholder_occurrences,
    _match_key,
    _para_full_text,
    _raw_placeholder_matches_from_doc,
)


def paragraph_dump(para):
    text = _para_full_text(para)
    return {
        "text_repr": repr(text),
        "runs": [repr(run.text) for run in para.runs],
        "matches": [match.group(0) for match in _PLACEHOLDER_RE.finditer(text)],
    }


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: python scripts/debug_docx_placeholders.py TEMPLATE.docx [PlaceholderName ...]", file=sys.stderr)
        return 2

    docx_path = Path(sys.argv[1])
    placeholders = sys.argv[2:]
    doc = Document(docx_path)

    print(f"placeholder_regex={_PLACEHOLDER_RE.pattern!r}")
    print(f"body_paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    for pi, para in enumerate(doc.paragraphs):
        dump = paragraph_dump(para)
        if dump["matches"] or any(token in dump["text_repr"] for token in placeholders):
            print(f"paragraph[{pi}] {json.dumps(dump, ensure_ascii=False)}")

    for ti, table in enumerate(doc.tables):
        print(f"table[{ti}] rows={len(table.rows)}")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    dump = paragraph_dump(para)
                    if dump["matches"] or any(token in dump["text_repr"] for token in placeholders):
                        print(f"table[{ti}].row[{ri}].cell[{ci}].paragraph[{pi}] {json.dumps(dump, ensure_ascii=False)}")

    keys = {_match_key(match) for match in _PLACEHOLDER_RE.finditer(_document_placeholder_scan_text(doc))}
    if placeholders:
        keys.update(placeholders)
    values = {key: key for key in keys}
    raw_matches = _raw_placeholder_matches_from_doc(doc, values)
    occurrences = _extract_placeholder_occurrences(doc, values)
    print("counts", json.dumps({
        "full_text_regex_matches": sum(1 for match in _PLACEHOLDER_RE.finditer(_document_placeholder_scan_text(doc)) if _match_key(match) in values),
        "source_aware_raw_matches": len(raw_matches),
        "occurrences": len(occurrences),
    }, ensure_ascii=False))
    print("raw_matches", json.dumps(raw_matches, ensure_ascii=False, indent=2))
    print("occurrences", json.dumps(occurrences, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
